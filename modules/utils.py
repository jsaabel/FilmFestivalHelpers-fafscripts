# 30.08.2021
from docx import Document
from modules import notion as n
import urllib.request
import datetime
from dateutil.parser import parse  # pip install python-dateutil


# # decorator functions
# def decorator_general(f):
#     """info decorator for all functions. Prints start message."""
#     def wrapper(*args, **kwargs):
#         print(f"\tExecuting {f}... ")
#         rv = f(*args, **kwargs)
#         return rv
#     return wrapper
#
#
# def decorator_get_functions(f):
#     """info decorator for functions that return a value."""
#     def wrapper(*args, **kwargs):
#         print(f"\tExecuting {f} with arg(s) {args}... ")
#         rv = f(*args, **kwargs)
#         print(f"\t\treturned {rv}.")
#         return rv
#     return wrapper


def add_commissioned_films(document):
    """adds formatted list of commissioned films to a document (returns document object)"""
    # 24.07.2021
    # integrate this into catalogue export?

    # FETCH FILTERED AND SORTED FILM DATABASE
    print('Trying to retrieve film database...')
    data_dict = dict()
    n.add_filter_to_request_dict(data_dict, 'Programme', 'multi_select', 'contains',
                                 'NBC: Commissioned Film')
    n.add_sorts_to_request_dict(data_dict, 'Seq', 'ascending')
    films = n.get_db('films', data_dict=data_dict)
    item_count_notion = n.get_item_count(films)

    # add film information

    for index in range(0, item_count_notion):
        film_title = n.get_property(index, 'English Title', 'title', source=films)
        if n.get_property(index, 'Original Title', 'text', source=films):
            film_title_ov = n.get_property(index, 'Original Title', 'text', source=films)
        else:
            film_title_ov = ""
        year = n.get_property(index, 'Year', 'select', source=films)
        director = n.get_property(index, 'Director', 'text', source=films)
        runtime = n.get_property(index, 'Runtime', 'text', source=films)
        country = list_to_comma_separated(n.get_property(index, 'Country', 'multi_select', source=films))

        paragraph = document.add_paragraph()
        run = paragraph.add_run(film_title)
        run.italic = True
        if film_title_ov:
            paragraph.add_run(f" / {film_title_ov}")
            run.italic = True
        paragraph.add_run(f" ({year}). {director}, {runtime}, {country}.")

    return document


def add_text(text_source: str, document_object):
    """add raw text from db into formatted word segment (synopsis/bio)"""
    # dealing with paragraphs and italics and preserving batch import compatibility
    try:
        textsrc_fixed = text_source.replace("</i>", "<i>").replace("<b>", "").replace("</b>", "")
        sections = textsrc_fixed.split("<br>")
    except AttributeError:
        pass
    else:
        for section in sections:
            paragraph = document_object.add_paragraph()
            sub_sections = section.split("<i>")
            if len(sub_sections) <= 1:
                run = paragraph.add_run(section)
            else:
                sub_section_count = 1
                for sub_section in sub_sections:
                    run = paragraph.add_run(sub_section)
                    if sub_section_count % 2 == 0:
                        run.italic = True
                    sub_section_count += 1
    return document_object


def catalogue_text_to_web_text(text_source: str, document_object):
    "replaces html tags for catalogue with tags for web"
    # text_2 = text.replace("<i>", "<em>")
    # text_3 = text_2.replace("</i>", "</em>")
    # text_4 = text_3.replace("<br>", "")
    # return text_4
    try:
        textsrc_fixed = text_source.replace("</i>", "</em>").replace("<i>", "<em>")
        sections = textsrc_fixed.split("<br>")
    except AttributeError:
        pass
    else:
        for section in sections:
            paragraph = document_object.add_paragraph(section)
    return document_object


def comma_separated_to_list(string: str):
    """Converts a comma separated string to a list object"""
    return string.split(', ')


def generate_film_files(data_source, folder_location_text: str = None, folder_location_images: str = None):
    """generates film documents and/or images from a provided notion database json object and saves them
    to specified folder(s) in order of database sort."""
    item_count_notion = n.get_item_count(data_source)

    # assign data
    for index in range(item_count_notion):
        # NB: make sure property names and types match those from database
        film_data = {
            "title": n.get_property(index, "English Title", 'title', source=data_source),
            "title_ov": n.get_property(index, "Original Title", 'text', source=data_source),
            "year": n.get_property(index, "Year", 'select', source=data_source),
            "director": n.get_property(index, "Director", 'text', source=data_source),
            "synopsis": n.get_property(index, "Synopsis", 'text', source=data_source),
            "bio": n.get_property(index, "Bio", 'text', source=data_source),
            "country": list_to_comma_separated(n.get_property(index, "Country", 'multi_select', source=data_source)),
            "runtime": n.get_property(index, "Runtime", 'text', source=data_source),  # OBS
            "technique": list_to_comma_separated(n.get_property(index, "Technique", 'multi_select', source=data_source)),
            "production": n.get_property(index, "Production", 'text', source=data_source),
            "animation": n.get_property(index, "Animation", 'text', source=data_source),
            "seq": str(index + 2).zfill(2),  # OBS
            "still_img": n.get_property(index, "Pic", 'url', source=data_source),
            # OBS 'new' functionality for feature films
            "programme": list_to_comma_separated(n.get_property(index, "Programme", 'multi_select', source=data_source)).lower(),
            "event_relations": n.get_property(index, 'Events', 'relation', source=data_source)
        }

        # initialize word document
        doc = Document()

        # add title and director paragraphs
        paragraph = doc.add_paragraph(film_data.get('title'))
        if film_data.get('title_ov'):
            paragraph.add_run(f" / {film_data.get('title_ov')}")
        paragraph.add_run(f" ({film_data.get('year')})")
        doc.add_paragraph(film_data['director'])

        # add info lines for feature films
        if 'feature' in film_data.get('programme'):
            if film_data['event_relations']:
                info_lines = n.generate_info_lines(film_data['event_relations'])
                for info_line in info_lines:
                    doc.add_paragraph(info_line)


        # add synopsis
        add_text(film_data.get('synopsis'), doc)

        # add bio
        add_text(film_data.get('bio'), doc)

        # add credits
        doc.add_paragraph(f"Country: {film_data.get('country')}")
        doc.add_paragraph(f"Runtime: {film_data.get('runtime')}")
        if film_data.get('technique'):
            doc.add_paragraph(f"Technique: {film_data.get('technique')}")
        if film_data.get('production'):
            doc.add_paragraph(f"Production: {film_data.get('production')}")
        if film_data.get('animation'):
            doc.add_paragraph(f"Animation: {film_data.get('animation')}")

        # avoid file name errors
        file_name = safe_file_name(f"{film_data['seq']} {film_data.get('title')}")

        # save word document
        if folder_location_text:
            doc_filename = f"{folder_location_text}/{file_name}.docx"
            doc.save(doc_filename)
            print(f"\tsaved word document for {film_data.get('title')}.")

        # save image
        if folder_location_images:
            if film_data.get("still_img"):
                save_img(film_data.get('still_img'), img_name=file_name, location=folder_location_images)


def get_date_string(notion_date, output):
    """parse date from notion, return a formatted string for day or time."""
    dt_object = parse(notion_date)
    date_string = str()
    if output == 'day':
        date_string = dt_object.strftime("%a, %d %b") # e.g.: Tue, 08 Jun
    elif output == 'time':
        date_string = dt_object.strftime("%H:%M")  # e.g.: 13:00
    return date_string


def get_tag_dict() -> dict:
    """return a dictionary matching eventive tag names to their eventive ids."""
    tag_to_tag_id = dict()
    tags = n.get_db('tags')
    tag_count = n.get_item_count(tags)
    for index in range(tag_count):
        tag_name = n.get_property(index, 'Name', 'title', source=tags)
        tag_id = n.get_property(index, 'eventive_id', 'text', source=tags)
        tag_to_tag_id[tag_name] = tag_id
    return tag_to_tag_id


def get_tag_dict_test() -> dict:
    """return a dictionary matching eventive tag names to their eventive test bucket ids."""
    tag_to_tag_id = dict()
    tags = n.get_db('tags')
    tag_count = n.get_item_count(tags)
    for index in range(tag_count):
        tag_name = n.get_property(index, 'Name', 'title', source=tags)
        tag_id = n.get_property(index, 'id_in_test_bucket', 'text', source=tags)
        tag_to_tag_id[tag_name] = tag_id
    return tag_to_tag_id


def list_to_comma_separated(list: list):
    """Converts a list object to a comma separated string"""
    joined_string = ", ".join(list)
    return joined_string


def print_error_message(function_name: str, error_details: str = None):
    print(f"\tError encountered while executing {function_name}.")
    if error_details:
        print(f"\t{error_details}")


def safe_file_name(text: str) -> str:
    """Takes a string and removes all characters that are invalid in filenames"""
    invalid_file_characters = ["/", ":", "*", "?", "<", ">", "|"]
    for invalid_file_character in invalid_file_characters:
        if invalid_file_character in text:
            text = text.replace(invalid_file_character, "")
    return text


def save_first_img(file_list: list, file_name: str):
    """download and save the first image in a list of files (from notion 'files' property) under the specified name."""
    try:
        if ".png" in file_list[0]:
            file_name = f"{file_name}.png"
        else:
            file_name = f"{file_name}.jpg"
        urllib.request.urlretrieve(file_list[0], file_name)
        print(f'\t... saved image: {file_name}.')
    except (IndexError, TypeError):  # avoid error message in case of missing file
        pass


def safe_folder_name(text: str) -> str:
    """Takes a string and removes all characters that are invalid in folder names"""
    invalid_file_characters = [":", "*", "?", "<", ">", "|"]
    for invalid_file_character in invalid_file_characters:
        if invalid_file_character in text:
            text = text.replace(invalid_file_character, "")
    return text


def save_img(img_url, img_name, location):
    """save image in chosen location in correct file format (.jpg/.png) and returns complete file path"""
    if ".png" in img_url:
        filename = f"{location}/{img_name}.png"
    else:
        filename = f"{location}/{img_name}.jpg"
    try:
        urllib.request.urlretrieve(img_url, filename)
        print(f"\t...saved image for: {img_name}.")
    except ValueError:
        pass


# MISC

# TAG EXTRACTOR
# while True:
#     url = input("Enter url: ")
#     # https://admin.eventive.org/event_buckets/60e7579087291c0083e504eb/tags/61000208cd49da003e11644b
#     slash = url.rfind("/")
#     print(url[slash + 1:])