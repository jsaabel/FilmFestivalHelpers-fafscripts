# 30.08.2021
from docx import Document
# from fafscripts.modules import notion_new as nn
import requests
import base64
import json
import urllib.request
import urllib.error
from dateutil.parser import parse  # pip install python-dateutil
import dropbox
import os
from fafscripts.scripts.forms import ProgrammeChoiceForm
import logging
import mammoth

logger = logging.getLogger(__name__)


def add_to_results(results: dict, title, category, message, url):

    code_to_category = {
        "e": "errors",
        "w": "warnings",
        "m": "messages",
    }

    feedback_dict = {
        "message": f"{title}: {message}",
        "url": url,
    }
    results[code_to_category[category]].append(feedback_dict)


def add_text(text_source: dict, document_object):
    """..."""
    try:
        paragraph = document_object.add_paragraph()
        for segment in text_source:
            run = paragraph.add_run(segment['text']['content'])
            if segment['annotations']['bold']:
                run.bold = True
            if segment['annotations']['italic']:
                run.italic = True
    except TypeError:
        logger.warning("NB: missing text encountered.")

    return document_object


def catalogue_text_to_web_text(text_source: str, document_object):
    "replaces html tags for catalogue with tags for web"
    # text_2 = text.replace("<i>", "<em>")
    # text_3 = text_2.replace("</i>", "</em>")
    # text_4 = text_3.replace("<br>", "")
    # return text_4
    try:
        textsrc_fixed = text_source.replace(
            "</i>", "</em>").replace("<i>", "<em>")
        sections = textsrc_fixed.split("<br>")
    except AttributeError:
        pass
    else:
        for section in sections:
            paragraph = document_object.add_paragraph(section)
    return document_object


def check_allcaps(text: str):
    for word in text.split():
        if word.isupper() and len(word.replace('.', '')) > 1:
            return True
    return False


def comma_separated_to_list(string: str):
    """Converts a comma separated string to a list object"""
    if not string:
        return None
    return string.split(', ')


def docx_to_html(docx_file):
    custom_styles = """ b => b.mark
                        u => u.initialism
                        p[style-name='Headline 1'] => h1
                        p[style-name='Text Margins'] => em 
                        p[style-name='Day'] => h2 
                        p[style-name='Event Time + Title'] => h3 
                        table => table.table.table-hover.striped
                        """

    def ignore_image(image):
        return []

    with open(docx_file, "rb") as f:
        result = mammoth.convert_to_html(
            f, style_map=custom_styles, convert_image=ignore_image)
        html = result.value

    os.remove(docx_file)
    return html


def download_file(file_url: str, target_folder: str, target_filename: str = ""):
    """download a file from provided/previously retrieved url."""
    # add file extension recognition?
    if not target_filename:
        parts = file_url.split('?')
        clean_url = parts[0]
        file_name_index = clean_url.rfind('/')
        target_filename = clean_url[file_name_index + 1:]
    logger.info(f'downloading file: {target_filename}...')
    urllib.request.urlretrieve(
        url=file_url, filename=f'{target_folder}/{target_filename}')


def ends_on_standard_character(text: str) -> bool:
    ending_characters = ['.', '!', '?']
    return True if text[-1] in ending_characters else False


def ends_with_whitespace(text: str) -> bool:
    return True if text[-1] == ' ' else False


def english_country_to_norwegian(country: str) -> str:
    "tries to find the Norwegian translation for a given country"

    nor_to_eng = {
        "Austria": "Østerrike",
        "Norway": "Norge",
        "Sweden": "Sverige",
        "Denmark": "Danmark",
        "Iceland": "Island",
        "Lithuania": "Litauen",
        "Estonia": "Estland",
        "France": "Frankrike",
        "Germany": "Tyskland",
        "Canada": "Kanada",
        "Spain": "Spania",
        "Uk": "Storbritannia",
    }

    return nor_to_eng.get(country, country)


def get_date_string(notion_date, output):
    """parse date from notion, return a formatted string for day or time."""
    dt_object = parse(notion_date)
    date_string = str()
    if output == 'day':
        date_string = dt_object.strftime("%a, %d %b")  # e.g.: Tue, 08 Jun
    elif output == 'time':
        date_string = dt_object.strftime("%H:%M")  # e.g.: 13:00
    return date_string


def get_results_dict():

    r = {
        'messages': [],
        'warnings': [],
        'errors': [],
    }

    return r


def list_to_comma_separated(list: list):
    """Converts a list object to a comma separated string"""
    joined_string = None
    try:
        joined_string = ", ".join(list)
    except TypeError:
        logger.warning("Could not convert to comma seperated: Empty field?")

    return joined_string


def print_error_message(function_name: str, error_details: str = None):
    logger.error(f"Error encountered while executing {function_name}.")
    if error_details:
        logger.error(f"{error_details}")


def safe_file_name(text: str) -> str:
    """Takes a string and removes all characters that are invalid in filenames"""
    invalid_file_characters = ["/", ":", "*", "?", "<", ">", "|"]
    for invalid_file_character in invalid_file_characters:
        if invalid_file_character in text:
            text = text.replace(invalid_file_character, "")
    return text


def save_first_img(file_list: list, file_name: str):
    """download and save the first image in a list of files (from notion 
    'files' property) under the specified name, return filename with
    appropriate file extension"""
    try:
        if ".png" in file_list[0]:
            file_name = f"{file_name}.png"
        else:
            file_name = f"{file_name}.jpg"
        urllib.request.urlretrieve(file_list[0], file_name)
        logger.info(f'... saved image: {file_name}.')
        return file_name
    except (IndexError, TypeError):  # avoid error message in case of missing file
        pass
    except urllib.error.HTTPError:
        logger.error(f'Failed to download {file_name}.')


def safe_folder_name(text: str) -> str:
    """Takes a string and removes all characters that are invalid in folder names"""
    invalid_file_characters = [":", "*", "?", "<", ">", "|"]
    for invalid_file_character in invalid_file_characters:
        if invalid_file_character in text:
            text = text.replace(invalid_file_character, "")
    return text


# NB: What about '&' etc? '&amp;'
# try if special letters work as intended, otherwise consult w3 tut on html entities
def rich_text_to_html(rich_text: list) -> str:
    res = ""
    if not rich_text:
        return None
    for segment in rich_text:
        text = segment['text']['content']
        # text = text.replace('\n', '<br>')
        if segment['annotations']['bold']:
            res += f"<strong>{text}</strong>"
        elif segment['annotations']['italic']:
            res += f"<em>{text}</em>"
        else:
            res += text
    paragraphs = res.split("\n")
    final = ""
    for paragraph in paragraphs:
        final += f"<p>{paragraph}</p>"
    return final


def rich_text_to_wordpress(rich_text: dict) -> str:
    res = ""
    for segment in rich_text:
        text = segment['text']['content']
        if segment['annotations']['bold']:
            res += f"<strong>{text}</strong>"
        elif segment['annotations']['italic']:
            res += f"<em>{text}</em>"
        else:
            res += text
    return res


def catalogue_text_to_web_text(text_source: str, document_object):
    "replaces html tags for catalogue with tags for web"
    # text_2 = text.replace("<i>", "<em>")
    # text_3 = text_2.replace("</i>", "</em>")
    # text_4 = text_3.replace("<br>", "")
    # return text_4
    try:
        textsrc_fixed = text_source.replace(
            "</i>", "</em>").replace("<i>", "<em>")
        sections = textsrc_fixed.split("<br>")
    except AttributeError:
        pass
    else:
        for section in sections:
            paragraph = document_object.add_paragraph(section)
    return document_object


def get_secret(name):
    with open('/etc/config.json') as config_file:
        config = json.load(config_file)
    return config.get(name)


def get_dbx():
    dbx = dropbox.Dropbox(app_key=get_secret("DROPBOX_APP_KEY"),
                          app_secret=get_secret("DROPBOX_APP_SECRET"),
                          oauth2_refresh_token=get_secret("DROPBOX_REFRESH_TOKEN"))
    return dbx


def dropbox_upload_from_url(url, target_location):
    try:
        response = urllib.request.urlopen(url)
    except urllib.error.HTTPError:
        logger.error(f"Encountered HTTP error when trying to download {url}."
                     + "Check if url can be opened.")
        return False
    data = response.read()
    dbx = get_dbx()
    try:
        dbx.files_upload(data, f"/{target_location}",
                         mode=dropbox.files.WriteMode.overwrite)
        logger.info(f"uploaded to dropbox: {target_location}.")
        return True
    except (dropbox.exceptions.ApiError, dropbox.exceptions.AuthError):
        logger.info(f"FAILED to upload {url} to dropbox :(")
        return False


def dropbox_upload_local_file(source_file, target_location, delete_on_success=True):

    if ":" in target_location:
        target_location = target_location.replace(":", "")
    dbx = get_dbx()
    with open(source_file, 'rb') as f:
        try:
            dbx.files_upload(
                f.read(), f"/{target_location}", mode=dropbox.files.WriteMode.overwrite)
            logger.info(f"uploaded to dropbox: {target_location}.")
            if delete_on_success:
                os.remove(source_file)
            return True
        except (dropbox.exceptions.ApiError, dropbox.exceptions.AuthError):
            logger.error(f"FAILED to upload {source_file} to dropbox :(")
            return False


# def handle_programme_choices(form: ProgrammeChoiceForm):
#     if not 'programme_choices' in globals():
#         global programme_choices
#         global seq_choices
#         programme_choices, seq_choices = n.generate_programmes_and_seq()
#     form.programme.choices = programme_choices
#     form.seq.choices = seq_choices
#     return form


def list_to_comma_separated(list: list):
    """Converts a list object to a comma separated string"""
    joined_string = ", ".join(list)
    return joined_string


def safe_folder_name(text: str) -> str:
    """Takes a string and removes all characters that are invalid in folder names"""
    invalid_file_characters = [":", "*", "?", "<", ">", "|"]
    for invalid_file_character in invalid_file_characters:
        if invalid_file_character in text:
            text = text.replace(invalid_file_character, "")
    return text


def save_img(img_url, img_name, location):
    """save image in chosen location in correct file format (.jpg/.png) and returns complete file path"""
    if ".png" in img_url:
        filename = f"{location}/{img_name}.png"
    else:
        filename = f"{location}/{img_name}.jpg"
    try:
        urllib.request.urlretrieve(img_url, filename)
        logger.info(f"...saved image for: {img_name}.")
    except ValueError:
        pass
    except urllib.error.HTTPError:
        logger.warning(f"failed to download from {img_url}.")


def wp_id_from_url(url: str):
    """uploads img in url to wp (if not already uploaded) and returns post id"""
    # handle dropbox links
    if "?dl=0" in url:
        url = url.replace("?dl=0", "?dl=1")
        filename = url.split('/')[-1][:-5]
    else:
        filename = url.split('/')[-1]
    logger.info(f"Trying to upload {filename} to wordpress...")

    if get_secret('FESTIVAL_URL') in url:
        logger.info(f"{filename} already exists, retrieving id only...")
        wp_url = get_secret('WP_BASE_URL') + "media?search=" + filename
        response = requests.get(wp_url)
        id = response.json()[0]['id']

    else:
        wp_url = get_secret("WP_BASE_URL") + "media"
        # wp_url = secrets.wp_base_url + config.get('')
        user = get_secret("WP_USER")
        password = get_secret("WP_API_KEY")
        credentials = user + ':' + password
        token = base64.b64encode(credentials.encode())
        content_disposition = "attachment; filename=\"" + filename + "\""
        headers = {
            'Authorization': 'Basic ' + token.decode("utf-8"),
            'Content-Disposition': content_disposition,
            'Content-Type': 'image/jpeg'
        }

        tempfile, _ = urllib.request.urlretrieve(url)
        payload = open(tempfile, 'rb')

        response = requests.post(wp_url, headers=headers, data=payload)

        id = response.json()['id']

        logger.info(f"uploaded {filename} to wordpress (id {id}).")

    return id
