from fafscripts.modules import notion as n, notion_new as nn, utils as u
from docx import Document
from random import randrange
import logging


def main(folder_location):
    # retrieve filtered and sorted location db
    logger = logging.getLogger(__name__)
    logger.info(f"Starting {__name__}")
    data_dict = dict()
    nn.add_filter_to_request_dict(
        data_dict, "Seq", "number", "is_not_empty", True)
    nn.add_sorts_to_request_dict(data_dict, 'Seq')
    notion_venues = nn.get_db("venues", data_dict=data_dict)
    # notion_venues_count = n.get_item_count(notion_venues)
    # compose document
    doc = Document()
    doc.add_paragraph("Venues")
    doc.add_paragraph()
    for venue_json in notion_venues['results']:
        v = nn.Page(json_obj=venue_json)
        venue_name = v.get_text('name')
        venue_text = v.get_text('catalogue-text')
        doc.add_paragraph(venue_name)
        doc.add_paragraph(venue_text)
        doc.add_paragraph()
    # save document
    rand = randrange(1000)
    doc.save(f'temp{rand}.docx')
    u.dropbox_upload_local_file(
        f"temp{rand}.docx", f"{folder_location}/venues.docx")
    logger.info(f'...saved venues.docx.')
