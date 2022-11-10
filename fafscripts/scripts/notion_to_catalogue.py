from fafscripts.modules import notion_new as n, utils as u, dbfuncs
from fafscripts.models import Venue, FilmProgramme
from docx import Document  # pip install python-docx
from fafscripts.scripts import (
    notion_to_catalogue_schedule, notion_to_catalogue_venues)
from random import randrange
import logging

logger = logging.getLogger(__name__)


def main(chosen_categories: list, text: bool, img: bool, chosen_programmes: list):

    logger.info(f"Starting notion_to_catalogue.py with arguments: {chosen_categories}, txt: {text}, img: {img}, "
                + f"chosen_programmes: {chosen_programmes}")
# GLOBAL VARIABLES
    dropbox_folder = u.get_secret("DROPBOX_FOLDER")
    catalogue_folder = f"{dropbox_folder}/Catalogue_22"  # cat. folder name
    filetypes = ["txt", "img"]  # sub-folders


# WORKING WITH PAGES
    logger.info("\nConnecting to notion database: Catalogue...")
    data_dict = dict()
    n.add_filter_to_request_dict(data_dict, 'data_source', 'select',
                                 'does_not_equal', '(nn)')
    n.add_sorts_to_request_dict(data_dict, 'Category 1')
    notion_data = n.get_db('catalogue', data_dict=data_dict)


# START LOOP
    for page_json in notion_data['results']:
        p = n.Page(json_obj=page_json)
        category_1 = p.get_text('category-1')
        if category_1 not in chosen_categories:  # skip if not in a chosen category
            continue
        category_2 = p.get_text('category-2')
        heading = p.get_text('heading')
        if not heading:
            logger.warning('Missing heading encountered. Skipping page.')
            continue
        logger.info(f"Working on {heading}...")
        data_source = p.get_text('data-source')
        if not data_source:
            logger.warning('No data source selected. Skipping page.')
            continue
        if data_source == "(nn)":
            continue
        seq = str(p.get_text('seq')).zfill(2)  # OBS str vs number etc..
        if not seq:
            logger.error(f"Missing sequence for {heading}.")

        event_relations = p.get_list('events')
        guest_relations = p.get_list('guests')
        film_relations = p.get_list('films')
        info_lines = None
        # if page is linked to an event in event db via relation property,
        # info lines are retrieved/generated from there.
        if event_relations:
            info_lines = generate_info_lines(event_relations)

        # setting folder/save locations and creating them
        if category_2:
            txt_location = f"{catalogue_folder}/{filetypes[0]}"\
                           f"/{category_1}/{category_2}"
            img_location = f"{catalogue_folder}/{filetypes[1]}"\
                           f"/{category_1}/{category_2}"
        else:
            txt_location = f"{catalogue_folder}/{filetypes[0]}/{category_1}"
            img_location = f"{catalogue_folder}/{filetypes[1]}/{category_1}"
        # making folder locations 'safe' by checking for unsupported characters etc
        txt_location = u.safe_folder_name(txt_location)
        img_location = u.safe_folder_name(img_location)

    # EXPORTS

        # dealing with different "data sources"/export modes

        if data_source == 'from file':

            if img:
                logger.debug("Data source: 'from file'")
                handle_file(p, img_location, seq, heading)

        elif data_source == 'schedule script':

            if text:
                logger.debug("Data source: 'schedule script'")
                # EXECUTE SCHEDULE GENERATOR SCRIPT
                notion_to_catalogue_schedule.main(folder_location=txt_location)

        elif data_source == 'venue script':

            if text:
                logger.debug("Data source: 'venue script'")
                notion_to_catalogue_venues.main(folder_location=txt_location)

        elif data_source == 'logo script':  # LOGO DOWNLOAD

            if img:
                logger.debug("Data source: 'logo script'")
                logo_main_folder = f'{img_location}/{seq} {heading}'
                logo_script(logo_main_folder)

        elif data_source == 'from page':

            if text:
                doc = Document()
                doc.add_paragraph(heading)
                if info_lines:
                    doc = add_info_lines(doc, info_lines)
                doc = add_text_from_page(doc, p, heading)
                save_and_upload_doc(doc, seq, heading, txt_location)

        # film programme texts
        elif data_source == 'txt from event' and event_relations:

            doc = Document()
            doc.add_paragraph(heading)
            if info_lines:
                doc = add_info_lines(doc, info_lines)
            e = n.Page(id=event_relations[0])
            if text:
                doc = add_event_text(doc, e)
                save_and_upload_doc(doc, seq, heading, txt_location)

        # guest page
        elif data_source == 'from db' and guest_relations:

            doc = Document()
            doc.add_paragraph(heading)
            g = n.Page(id=guest_relations[0])
            doc = add_biography(doc, g)
            if text:
                save_and_upload_doc(doc, seq, heading, txt_location)
            if img:
                pic = g.get_text('headshot-url')
                if pic:
                    if '?dl' in pic:
                        pic = pic.replace('?dl=0', '?dl=1')
                        u.dropbox_upload_from_url(
                            pic, target_location=f"{img_location}/{seq} {heading}.{pic.split('.')[-1][:-5]}")
                    else:
                        u.dropbox_upload_from_url(
                            pic, target_location=f"{img_location}/{seq} {heading}.{pic.split('.')[-1]}")
                else:
                    logger.warning(f'Missing image for {heading}.')

        # event page
        elif data_source == 'from db' and event_relations:

            doc = Document()
            doc.add_paragraph(heading)
            if info_lines:
                doc = add_info_lines(doc, info_lines)
            e = n.Page(id=event_relations[0])
            if text:
                doc = add_event_text(doc, e)
                save_and_upload_doc(doc, seq, heading, txt_location)
            if img:
                event_img = e.get_text('pic')
                if event_img:
                    if '?dl' in event_img:
                        event_img = event_img.replace('?dl=0', '?dl=1')
                        img_name = u.safe_file_name(
                            f'{seq} {heading}.{event_img.split(".")[-1][:-5]}')
                        u.dropbox_upload_from_url(
                            event_img, target_location=f"{img_location}/{img_name}")
                    # u.save_img(img_url=event_img, img_name=f'{seq} {heading}', location=img_location)
                    else:
                        img_name = f'{seq} {heading}.{event_img.split(".")[-1]}'
                        u.dropbox_upload_from_url(
                            event_img, f'{img_location}/{u.safe_file_name(img_name)}')
                else:
                    logger.warning(f'Missing image for {heading}.')

        # commissioned
        elif data_source == "commissioned script" and text:
            generate_commissioned_document(seq, heading, txt_location)

        # film page(s)
        elif data_source == 'from db' and film_relations:  # "regular"/non script-related export

            # analyze related film page and continue accordingly
            f = n.Page(id=film_relations[0])
            notion_programme_tag_id = f.get_list('film-programmes')[0]
            programme_query = FilmProgramme.query.filter_by(
                notion_id=notion_programme_tag_id).first()
            programme = programme_query.name
            seq = FilmProgramme.query.filter_by(
                notion_id=notion_programme_tag_id).first().name
            if chosen_programmes and programme not in chosen_programmes:
                continue
            logger.info(f"Working on {programme}.")
            print(f"NOT skipping {programme}!!")

            # (add Commissioned Films)
            # if notion_programme_tag_id == commissioned_programme_id:  # OBS
            #
            #     logger.info('Adding commissioned films...')
            #     doc = Document()
            #     doc.add_paragraph()  # blank line
            #     doc = add_commissioned_films(doc)  # special function for generating commissioned films document
            #     save_and_upload_doc(doc, seq, heading, txt_location)

            data_dict = dict()
            n.add_filter_to_request_dict(data_dict, '🎥 Film programmes', 'relation', 'contains',
                                         notion_programme_tag_id)
            n.add_sorts_to_request_dict(data_dict, programme_query.seq)
            films = n.get_db('films', data_dict=data_dict)
            generate_film_files(films, folder_location_text=txt_location,
                                folder_location_images=img_location, img=img, text=text)

        else:
            logger.warning(
                "Encountered unknown combination of data source and/or relation.")
            pass


def generate_info_lines(ids: list) -> list:
    """generates infolines for use in catalogue from a list of ids (typically from relation-property)"""
    logger.info("Generating info lines.")
    info_lines = list()
    try:
        for id in ids:

            # notion_data = get_page(id)
            event_page = n.Page(id=id)
            venue_id = event_page.get_list('venue')
            # venue_id = get_property_from_page('Venue', 'relation', source=notion_data)
            # venue_page = n.Page(id=venue_id) # this should probably done w/ model query...
            venue = dbfuncs.get_name_from_notion_id(venue_id[0], model=Venue)
            # venue = venue_page.get_text('name')
            day = event_page.get_date('time')[0][8:10]
            time = event_page.get_date('time')[0][11:16]
            age_limit = event_page.get_text('age-limit')
            info_line = f"{day} OCT / {time} / {venue.upper()}"
            if age_limit:
                info_line += f' / AGE LIMIT: {age_limit}'
            info_lines.append(info_line)
    except (IndexError, TypeError):
        logger.error(
            "Error encountered while generating info lines. Check that venue and time are set?")
        return []
    return sorted(info_lines)


def generate_commissioned_document(seq, heading, txt_location):
    """adds formatted list of commissioned films to a document and saves/uploads it"""

    document = Document()
    # FETCH FILTERED AND SORTED FILM DATABASE
    logger.info('Trying to retrieve film database...')
    data_dict = dict()
    # n.add_filter_to_request_dict(data_dict, 'Programme', 'multi_select', 'contains',
    #         'Nordic-Baltic Commissioned Films') # OBS: has to be updated for 2022!
    n.add_sorts_to_request_dict(data_dict, 'Seq', 'ascending')
    n.add_filter_to_request_dict(data_dict, '🎥 Film programmes', 'relation', 'contains',
                                 get_commissioned_programme_id())
    films = n.get_db('films', data_dict=data_dict)
    # item_count_notion = n.get_item_count(films)

    # add film information

    for film_json in films['results']:
        # for index in range(item_count_notion):
        f = n.Page(json_obj=film_json)
        film_title = f.get_text('english-title')
        if f.get_text('original-title'):
            film_title_ov = f.get_text('original-title')
        else:
            film_title_ov = ""
        year = f.get_text('year')
        director = f.get_text('director')
        runtime = f.get_text('runtime')
        country = f.get_text('country')

        paragraph = document.add_paragraph()
        run = paragraph.add_run(film_title)
        run.italic = True
        if film_title_ov:
            paragraph.add_run(f" / {film_title_ov}")
            run.italic = True
        paragraph.add_run(f" ({year}). {director}, {runtime}, {country}.")

    save_and_upload_doc(document, seq, heading, txt_location)


def generate_film_files(data_source, img: bool, text: bool, folder_location_text: str = None, folder_location_images: str = None):
    """generates film documents and/or images from a provided notion database json object and saves them
    to specified folder(s) in order of database sort."""

    # assign data
    for seq, film_json in enumerate(data_source['results']):
        # NB: make sure property names and types match those from database
        f = n.Page(json_obj=film_json)

        film_data = f.get_plain_text_dict()
        film_data['bio'] = f.get_rich_text('bio')
        film_data['synopsis'] = f.get_rich_text('synopsis')
        film_data['seq'] = str(seq + 2).zfill(2)
        # matching programme ids to programme titles
        programme_ids = f.get_list('film-programmes')
        programmes = [FilmProgramme.query.filter_by(
            notion_id=id).first().name for id in programme_ids]
        film_data['programme'] = u.list_to_comma_separated(programmes)
        file_name = u.safe_file_name(
            f"{film_data['seq']} {film_data.get('english-title')}")

        # initialize word document
        if text:
            doc = Document()

            # digital approval
            if film_data.get('digital-approval'):
                doc.add_paragraph("[---ON DIGITAL FESTIVAL---]")

            # add title and director paragraphs
            paragraph = doc.add_paragraph(film_data.get('english-title'))
            if film_data.get('original-title'):
                paragraph.add_run(f" / {film_data.get('original-title')}")
            paragraph.add_run(f" ({film_data.get('year')})")
            doc.add_paragraph(film_data['director'])

            # add info lines for feature films
            if 'feature' in film_data.get('programme').lower():
                if film_data['events']:  # OBS
                    info_lines = generate_info_lines(
                        u.comma_separated_to_list(film_data['events']))  # OBS
                    for info_line in info_lines:
                        doc.add_paragraph(info_line)

            # add synopsis
            u.add_text(film_data.get('synopsis'), doc)

            # add bio
            u.add_text(film_data.get('bio'), doc)

            # add credits
            doc.add_paragraph(f"Country: {film_data.get('country')}")
            doc.add_paragraph(f"Runtime: {film_data.get('runtime')}")
            if film_data.get('school'):
                doc.add_paragraph(f"School: {film_data.get('school')}")
            if film_data.get('technique'):
                doc.add_paragraph(f"Technique: {film_data.get('technique')}")
            if film_data.get('production'):
                doc.add_paragraph(f"Production: {film_data.get('production')}")
            if film_data.get('animation'):
                doc.add_paragraph(f"Animation: {film_data.get('animation')}")

            # save word document
            if folder_location_text:
                save_and_upload_doc(doc, film_data['seq'], heading=film_data.get(
                    'english-title'), txt_location=folder_location_text)

        # save image
        if img and folder_location_images and film_data.get('pic'):
            file_name = u.safe_file_name(
                f"{film_data['seq']} {film_data.get('english-title')}")
            # save_img(film_data.get('still_img'), img_name=file_name, location=folder_location_images)
            u.dropbox_upload_from_url(film_data.get('pic'),
                                      target_location=f"{folder_location_images}/{file_name}.{film_data.get('pic').split('.')[-1]}")


def add_biography(doc: Document, g: n.Page) -> Document:

    guest_text = g.get_rich_text('bio-eng')
    u.add_text(guest_text, doc)
    return doc


def add_info_lines(doc: Document, info_lines: list) -> Document:
    for info_line in info_lines:
        doc.add_paragraph(info_line)
    return doc


def add_event_text(doc: Document, e: n.Page) -> Document:
    event_text = e.get_rich_text('text-catalogue')
    u.add_text(event_text, doc)
    return doc
    # event_img = event_page.get_text('pic')


def add_text_from_page(doc: Document, p: n.Page, heading: str) -> Document:

    p.retrieve_children()
    text_blocks = p.get_children()
    # block_count = n.get_item_count(text_blocks)
    if not text_blocks:
        logger.warning(f"No text found on page '{heading}'.")
        return doc
    for text_block in text_blocks:  # could/should be made into function...
        try:
            property_type = text_block['type']
            rich_text = text_block[property_type]['rich_text']
            u.add_text(rich_text, doc)
        except IndexError:  # this is necessary to deal with, amongst others, blank lines.
            # OBS: Check if this still applies / has correct exception type!
            doc.add_paragraph()
    return doc


def get_commissioned_programme_id():

    return "9640db1a0b8b443ca58210baa5d94189"


def handle_file(p: n.Page, img_location: str, seq: str, heading: str):

    dl_urls = p.get_list('file')
    if not dl_urls:
        logger.error(
            f'OBS: Encountered missing file when trying to download {heading}')
        pass
    else:
        if ".png" in dl_urls[0]:
            file_extension = "png"
        elif ".pdf" in dl_urls[0]:
            file_extension = "pdf"
        elif ".zip" in dl_urls[0]:
            file_extension = "zip"

        else:
            file_extension = "jpg"
        file_name = u.safe_file_name(f'{seq} {heading}.{file_extension}')
        rand = randrange(1000)
        try:
            u.download_file(dl_urls[0], target_folder=".",
                            target_filename=f'temp{rand}.{file_extension}')
            u.dropbox_upload_local_file(f"temp{rand}.{file_extension}",
                                        f"{img_location}/{file_name}")
        except (IndexError, UnboundLocalError):
            logger.error(
                f'Encountered error when trying to download {heading}!')


def logo_script(logo_main_folder):

    request_dict = dict()
    n.add_filter_to_request_dict(request_dict, property_name='Catalogue',
                                 property_type='checkbox', filter_condition='equals',
                                 filter_content=True)
    notion_supporters = n.get_db('supporters', data_dict=request_dict)
    for supporter_json in notion_supporters['results']:
        # gather data
        supporter = n.Page(json_obj=supporter_json)
        name = supporter.get_text('name')
        logger.info(f"Trying to export logos for {name}.")
        category = supporter.get_text('category')
        seq = supporter.get_text('seq')
        logos_urls = supporter.get_list('logos')
        if not (category and seq and logos_urls):
            logger.error(
                f"Missing required information for supporter {name}. Skipping.")
            continue
        # create folders
        logo_sub_folder_location = f'{logo_main_folder}/{category}/'\
                                   f'{str(seq).zfill(2)} {name}'
        for logo_url in logos_urls:
            filename = logo_url.split('?')[0].split('/')[-1]
            u.dropbox_upload_from_url(logo_url,
                                      target_location=f"{logo_sub_folder_location}/{filename}")


def save_and_upload_doc(doc: Document, seq, heading, txt_location):

    doc_filename = u.safe_file_name(
        f"{seq} {heading}.docx")  # making filename 'safe'
    rand = randrange(1000)
    doc.save(f"temp{rand}.docx")
    u.dropbox_upload_local_file(
        f"temp{rand}.docx", f"{txt_location}/{doc_filename}")


# def save_guest_img(g: n.Page, img_path: str) -> None:
#
#     rand = randrange(1000)
#     temp = u.save_first_img(guest_img, f'temp{rand}')  # replace with n.download_file?
#     u.dropbox_upload_local_file(temp, target_location=img_path)
    # target_location=f"{img_location}/{seq} {heading}.{temp.split('.')[-1]}")
