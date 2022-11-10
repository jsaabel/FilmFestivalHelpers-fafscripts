# SCHEDULE EXPORT
from fafscripts.modules import notion_new as n, utils as u, notion as nl
from docx import Document
import datetime
from random import randrange
import logging
# from docx2pdf import convert

logger = logging.getLogger(__name__)


def get_nn_string():
    return "[not specified]"


def get_travel_info(guest: n.Page):
    """compile a dictionary with all relevant travel information"""
    # generate dictionary
    nn_string = get_nn_string()
    # new approach for this part 12.10.2021: Fix data / account for missing info here!
    travel_dict = dict()
    # nn_string = "[not specified]"

    travel_dict["incoming"] = guest.get_text('incoming')
    departure_date_notion = guest.get_date('departure-incoming')
    if departure_date_notion:
        travel_dict["departure_incoming_date"] = u.get_date_string(
            notion_date=departure_date_notion[0], output='day')
        travel_dict["departure_incoming_time"] = u.get_date_string(
            notion_date=departure_date_notion[0], output='time')
    else:
        travel_dict["departure_incoming_date"] = nn_string
        travel_dict["departure_incoming_time"] = nn_string
    travel_dict["booking_ref_in"] = guest.get_text('booking-ref-in')
    travel_dict["arrival_at"] = guest.get_text('arrival-at')
    arrival_date = guest.get_date('arrival')
    if arrival_date:
        travel_dict["arrival_date"] = u.get_date_string(
            notion_date=arrival_date[0], output='day')
        travel_dict["arrival_time"] = u.get_date_string(
            notion_date=arrival_date[0], output='time')
    else:
        travel_dict["arrival_date"] = nn_string
        travel_dict["arrival_time"] = nn_string
    travel_dict["transport_incoming"] = guest.get_text('transport-incoming')
    travel_dict["outgoing"] = guest.get_text('outgoing')
    travel_dict["booking_ref_out"] = guest.get_text('booking-ref-out')
    travel_dict["departure_outgoing_from"] = guest.get_text('departure-from')
    departure_outgoing_date_notion = guest.get_date('departure-outgoing')
    if departure_outgoing_date_notion:
        travel_dict["departure_outgoing_date"] = u.get_date_string(
            notion_date=departure_outgoing_date_notion[0], output='day')
        travel_dict["departure_outgoing_time"] = u.get_date_string(
            notion_date=departure_outgoing_date_notion[0], output='time')
    else:
        travel_dict["departure_outgoing_date"] = nn_string
        travel_dict["departure_outgoing_time"] = nn_string
    travel_dict["transport_outgoing"] = guest.get_text('transport-outgoing')

    for key, value in travel_dict.items():  # setting all missing dic values to nn_string
        if not value:
            travel_dict[key] = nn_string

    return travel_dict


def main(name: str, all: bool = False, pdf: bool = False, upload: bool = True):
    dropbox_folder = u.get_secret("DROPBOX_FOLDER")
    # location of our schedule template file
    template = "files/schedule_template.docx"
    filter_dict = dict()
    notion_data = n.get_db('guests', data_dict=filter_dict)
    data_dict = dict()
    # retrieved event db has to be sorted chronologically
    n.add_sorts_to_request_dict(data_dict, 'Time', 'ascending')
    events_data = n.get_db('events', data_dict=data_dict)
    item_count_events = nl.get_item_count(events_data)
    notion_venues = n.get_db('venues')

# looping over guests
    for guest_json in notion_data['results']:
        g = n.Page(json_obj=guest_json)
        guest_name = g.get_text('name')
        if guest_name:
            guest_name = guest_name.strip()
        if not all and guest_name != name:
            continue  # skipping all other guests if 'single mode' is selected
        # OBS THIS SHOULD BE IMPLEMENTED WITH FILTER INSTEAD!
        # check if assigned to events, skip guest if not // OBS: Is this even necessary if db is filtered upon retrieval?
        event_ids_attendee = g.get_list('events-attendees')
        event_ids_functionary = g.get_list('events-functionaries')
        if not (event_ids_attendee or event_ids_functionary):
            logger.debug(f"Skipping {guest_name}: No events assigned.")
            continue  # skipping guests if no events are assigned

        logger.info(f"Found {len(event_ids_attendee)} (Attendee) /"
                    + f" {len(event_ids_functionary)} (Functionary) events for {guest_name}.")
        guest_id = g.id
        # events are retrieved from event db using relation property
        event_ids_combined = list(event_ids_attendee)
        event_ids_combined.extend(
            x for x in event_ids_functionary if x not in event_ids_attendee)
        # composing document
        try:
            document = Document(template)  # loading in schedule template
        except ValueError:
            logger.error(f"Failed to open {template}, aborting script")
            raise FileNotFoundError
        # Title Paragraph
        paragraph = document.paragraphs[0]
        paragraph.add_run(f"SCHEDULE FOR {guest_name.upper()}")
        paragraph.style = "Headline 1"
        # As of ...
        today = datetime.datetime.now()
        paragraph = document.add_paragraph(f"(Information as of {today.strftime('%d')} {today.strftime('%b')}. "
                                           f"Changes and/or mistakes may occur. Please refer to your ticket(s) for further details.)")
        # make sure these paragraph styles are defined in the template
        paragraph.style = "Text Margins"
        # Blank line
        document.add_paragraph("")
        # OPTIONAL: "Intro" paragraph - instructions on eventive, covid etc
        first_name = guest_name.split(" ", 1)[0]
        paragraph = document.add_paragraph(f"Dear {first_name}! We are looking forward to welcoming you to Fredrikstad,"
                                           f" and hope you will find this schedule helpful. Please make sure to read the"
                                           f" description of your events carefully, as rehearsals, tech checks etc."
                                           f" are often scheduled BEFORE the official start time of the event.")
        paragraph.style = "Event Text"

        # TRAVEL INFORMATION (temporatily disabled for easier testing of other functionality)
        # check status of 'Travel?' checkbox
        travel_check = g.get_text('travel')  # OBS
        travel_info = dict()  # getting rid of 'may not be assigned' warning
        if travel_check:

            # return dictionary with travel information
            travel_info = get_travel_info(g)
            paragraph = document.add_paragraph("ARRIVAL")
            paragraph.style = "Day"

            paragraph = document.add_paragraph("Departure with ")
            paragraph.style = "Event Text"
            run = paragraph.add_run(f"{travel_info['incoming']} ")
            run.bold = True
            run = paragraph.add_run("at ")
            run.bold = False
            run = paragraph.add_run(
                f"{travel_info['departure_incoming_time']}")
            run.bold = True
            run = paragraph.add_run(" (local time) on ")
            run.bold = False
            run = paragraph.add_run(
                f"{travel_info['departure_incoming_date']}. ")
            run.bold = True

            if get_nn_string() not in travel_info['booking_ref_in']:
                run = paragraph.add_run(
                    f"Your booking reference: {travel_info['booking_ref_in']}. ")
                run.bold = False

            paragraph = document.add_paragraph(f"Arrival at ")
            paragraph.style = "Event Text"
            run = paragraph.add_run(f"{travel_info['arrival_at']}")
            run.bold = True
            run = paragraph.add_run(f" at ")
            run.bold = False
            run = paragraph.add_run(f"{travel_info['arrival_time']}. ")
            run.bold = True

            if '[' not in travel_info['transport_incoming']:
                run = paragraph.add_run(
                    f"Transport to festival: {travel_info['transport_incoming']}. ")
                run.bold = False

        # Split event ids into days
        event_ids_wednesday = list()
        event_ids_thursday = list()
        event_ids_friday = list()
        event_ids_saturday = list()
        event_ids_sunday = list()

        for event_id in event_ids_combined:
            for i2 in range(item_count_events):
                # matching event id obtained from relation to event db
                if event_id == events_data['results'][i2]['id']:
                    found_index = i2
                    # tuple: index in (sorted) json - event id. allows for ordering!
                    index_and_id = (found_index, event_id)

                    date = nl.get_property(
                        i2, 'Time', property_type='date', source=events_data)
                    if not date:
                        logger.error(
                            f"Encountered missing time for event id {event_id}! Skipping this event.")
                        continue
                    # matching dates and weekdays
                    if '10-19' in date[0]:
                        event_ids_wednesday.append(index_and_id)
                    elif '10-20' in date[0]:
                        event_ids_thursday.append(index_and_id)
                    elif '10-21' in date[0]:
                        event_ids_friday.append(index_and_id)
                    elif '10-22' in date[0]:
                        event_ids_saturday.append(index_and_id)
                    elif '10-23' in date[0]:
                        event_ids_sunday.append(index_and_id)

        weekday_dict = {
            'Wednesday, 19 OCT': sorted(event_ids_wednesday),
            'Thursday, 20 OCT': sorted(event_ids_thursday),
            'Friday, 21 OCT': sorted(event_ids_friday),
            'Saturday, 22 OCT': sorted(event_ids_saturday),
            'Sunday, 23 OCT': sorted(event_ids_sunday),
        }

        seen_days = 0
        for day, ids in weekday_dict.items():
            if len(ids) >= 1:  # -> only add this paragraph if at least one event for the day exists
                seen_days += 1
                # if seen_days > 1:
                #     document.add_page_break()
                paragraph = document.add_paragraph(day)
                paragraph.style = "Day"
                for id in ids:
                    for i3 in range(item_count_events):
                        if events_data['results'][i3]['id'] == id[1]:  # assign event info
                            event_title = nl.get_property(
                                i3, 'Name', property_type='title', source=events_data)
                            if ' (second screening)' in event_title:
                                event_title = event_title.replace(
                                    ' (second screening)', '')
                            event_time = nl.get_property(
                                i3, 'Time', property_type='date', source=events_data)
                            if not event_time:
                                logger.error(
                                    f"No time set for event {event_title}! Skipping this event...")
                                continue
                            venue_notion_id = nl.get_property(
                                i3, 'Venue', 'relation', source=events_data)
                            if not venue_notion_id:
                                logger.error(
                                    f"No venue set for event f{event_title}! Skipping this event...")
                                continue
                            venue_notion_page = nl.get_page_from_db(
                                venue_notion_id[0], source=notion_venues)
                            event_venue = nl.get_property_from_page(
                                'Name', 'title', source=venue_notion_page)
                            event_description = nl.get_property(
                                i3, 'Description Guest Schedule', source=events_data)
                            # event_attendees = n.get_property(i3, 'Attendees', property_type='relation',source=events_data)
                            event_functionaries = nl.get_property(
                                i3, 'Functionaries', property_type='relation', source=events_data)
                            event_description_functionaries = nl.get_property(
                                i3, 'Description Functionaries', source=events_data)

                            # write to document
                            paragraph = document.add_paragraph(f"{u.get_date_string(event_time[0], output='time')} "
                                                               f"{event_title} at {event_venue}")
                            paragraph.style = "Event Time + Title"
                            if guest_id in event_functionaries:
                                paragraph = document.add_paragraph(
                                    event_description_functionaries)
                            else:
                                paragraph = document.add_paragraph(
                                    event_description)
                            paragraph.style = "Event Text"

        # outgoing travel information
        if travel_check:
            if not g.get_text('departure-from'):
                pass
            else:
                paragraph = document.add_paragraph("DEPARTURE")
                paragraph.style = "Day"
                paragraph = document.add_paragraph("Departure with ")
                paragraph.style = "Event Text"
                run = paragraph.add_run(
                    f"{travel_info['outgoing']} from {travel_info['departure_outgoing_from']}")
                run.bold = True
                run = paragraph.add_run(" at ")
                run.bold = False
                run = paragraph.add_run(
                    f"{travel_info['departure_outgoing_time']}")
                run.bold = True
                run = paragraph.add_run(" on ")
                run.bold = False
                run = paragraph.add_run(
                    f"{travel_info['departure_outgoing_date']}. ")
                run.bold = True

                if '[' not in travel_info['booking_ref_out']:
                    run = paragraph.add_run(
                        f"Your booking reference: {travel_info['booking_ref_out']}. ")
                    run.bold = False

                if '[' not in travel_info['transport_outgoing']:
                    run = paragraph.add_run(
                        f"Transport: {travel_info['transport_outgoing']}. ")

        # Save Document
        filename = f"{u.get_secret('FESTIVAL_ACRONYM')}_Schedule_{guest_name}.docx"
        # rand = randrange(1000)
        document.save(f"schedule_temp_{g.id}.docx")

        logger.info(f"locally saved schedule for {guest_name}.")
        if not upload:
            return
        u.dropbox_upload_local_file(f"schedule_temp_{g.id}.docx",
                                    f"{dropbox_folder}/schedules/{filename}")
