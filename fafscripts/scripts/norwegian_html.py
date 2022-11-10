from fafscripts.modules import notion_new as n, utils as u, eventive, dbfuncs
from fafscripts.models import Venue
import docx
from docx import Document
from random import randrange
import logging


def generate_guests_document():
    logger = logging.getLogger(__name__)
    data_dict = dict()
    n.add_filter_to_request_dict(
        data_dict, 'Wordpress', 'url', 'is_not_empty', True)
    guests_db = n.get_db('guests', data_dict=data_dict)

    doc = Document()

    for guest_json in guests_db['results']:

        g = n.Page(guest_json)

        name = g.get_text('name')

        bio_norw = g.get_rich_text('bio-norw')
        if not bio_norw:
            logger.warning(
                f"No Norwegian text found for {name}. Skipping this guest.")
            continue

        doc.add_paragraph(name)
        p = doc.add_paragraph()
        add_hyperlink(p, g.get_text('wordpress'),
                      "(Click here to go to WordPress)")

        bio_norw = u.rich_text_to_html(bio_norw)

        doc.add_paragraph(bio_norw)
        doc.add_paragraph()

    rand = randrange(1000)

    dropbox_folder = u.get_secret("DROPBOX_FOLDER")
    doc.save(f"temp{rand}.docx")
    u.dropbox_upload_local_file(
        f'temp{rand}.docx', f"{dropbox_folder}/guests_norwegian_html.docx")


def generate_events_document():
    logger = logging.getLogger(__name__)
    data_dict = dict()
    n.add_filter_to_request_dict(
        data_dict, 'Wordpress', 'url', 'is_not_empty', True)
    events_db = n.get_db('events', data_dict=data_dict)

    doc = Document()

    for event_json in events_db['results']:

        e = n.Page(event_json)

        name = e.get_text("name")

        text_norw = e.get_rich_text('text-norw')
        venue = e.get_list('venue')
        time = e.get_date('time')

        if not all([text_norw, venue, time]):
            logger.warning(
                f"No Norwegian text/Venue/Time found for {name}. Skipping this event.")
            continue

        doc.add_paragraph(name)
        p = doc.add_paragraph()
        add_hyperlink(p, e.get_text('wordpress'),
                      "(Click here to go to WordPress)")

        venue = dbfuncs.get_name_from_notion_id(venue[0], Venue)
        doc.add_paragraph(f"Where: {venue}")
        date_string = f"{u.get_date_string(time[0], 'day')}, {u.get_date_string(time[0], 'time')}"
        doc.add_paragraph(f"When: {date_string}")

        text_norw = u.rich_text_to_html(text_norw)

        eventive_url = e.get_text('eventive')
        eventive_id = eventive.url_to_tag_id(
            eventive_url) if eventive_url else ""
        eventive_btn_code = ""
        if eventive_id:
            eventive_btn_code = f"<br><a class='btn-faf' target='_blank' href='https://{u.get_secret('FESTIVAL_ACRONYM').lower()}.eventive.org/schedule/{eventive_id}'>Billetter</a>"
            text_norw += eventive_btn_code

        doc.add_paragraph(text_norw)

    rand = randrange(1000)

    dropbox_folder = u.get_secret("DROPBOX_FOLDER")
    doc.save(f"temp{rand}.docx")
    u.dropbox_upload_local_file(
        f'temp{rand}.docx', f"{dropbox_folder}/events_norwegian_html.docx")


def main(guests: bool = True, events: bool = True):

    logger = logging.getLogger(__name__)
    if guests:
        logger.info("Trying to generate Norwegian html document for Guests.")

        generate_guests_document()

    if events:
        logger.info("Trying to generate Norwegian html document for Events.")
        generate_events_document()


# stolen from stackoverflow
def add_hyperlink(paragraph, url, text, color="FF8822", underline=True):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink
