from fafscripts.modules import notion_new as n, utils as u
from fafscripts.models import FilmProgramme
from docx import Document
from random import randrange
import logging


def main(programme, seq, norwegian_mode):

    logger = logging.getLogger(__name__)

    dropbox_folder = u.get_secret("DROPBOX_FOLDER")

# retrieving a filtered and sorted film db
    programme_id = FilmProgramme.query.filter_by(
        name=programme).first().notion_id
    data_dict = dict()
    n.add_filter_to_request_dict(data_dict, '🎥 Film programmes', 'relation',
                                 'contains', programme_id)
    n.add_sorts_to_request_dict(data_dict, seq)
# n.add_sorts_to_request_dict(data_dict, 'Seq', 'ascending') # sort by sequence (if exists)
    data_source = n.get_db('films', data_dict=data_dict)

    doc = Document()

    doc.add_paragraph(programme)

    for i, film_json in enumerate(data_source['results']):
        # NB: make sure property names and types match those from database
        f = n.Page(json_obj=film_json)
        countries = f.get_list('country')
        if norwegian_mode:
            countries = [u.english_country_to_norwegian(
                country) for country in countries]

        film_data = {
            "title": f.get_text('english-title'),
            # "title_ov": n.get_property(i, "Original Title", 'rich_text', source=data_source),
            # "year": n.get_property(i, "Year", 'select', source=data_source),
            "director": f.get_text('director'),
            # "synopsis": n.get_property(i, "Synopsis", 'rich_text', source=data_source),
            # "bio": n.get_property(i, "Bio", 'rich_text', source=data_source),
            "country": u.list_to_comma_separated(countries),
            # "runtime": n.get_property(i, "Runtime", 'rich_text', source=data_source),  # OBS
            # "technique": u.list_to_comma_separated(n.get_property(i, "Technique", 'multi_select', source=data_source)),
            # "production": n.get_property(i, "Production", 'rich_text', source=data_source),
            # "animation": n.get_property(i, "Animation", 'rich_text', source=data_source),
            # "seq": str(i + 2).zfill(2),  # OBS
        }

        # composing and adding a single line for each film in word document
        p = doc.add_paragraph()

        run = p.add_run(film_data.get('title'))
        run.italic = True

        run = p.add_run(f" – {film_data['director']} – {film_data['country']}")
        run.italic = False

    rand = randrange(1000)
    doc.save(f'temp{rand}.docx')
    logger.info('List written.')

    if norwegian_mode:
        u.dropbox_upload_local_file(
            f'temp{rand}.docx', f"{dropbox_folder}/film_lists/{programme}_norsk.docx")
    else:
        u.dropbox_upload_local_file(
            f'temp{rand}.docx', f"{dropbox_folder}/film_lists/{programme}.docx")
