import datetime
from docx import Document
from fafscripts.modules import notion_new as n, utils as u, dbfuncs
from fafscripts.models import Guest, Venue, Event
import logging


def main():

    # result dict
    r = u.get_results_dict()

    # retrieve db
    data = dict()
    n.add_sorts_to_request_dict(
        request_dict=data, property_name="Date", sorts_direction='ascending')
    tasks = n.get_db('volunteer_schedule', data_dict=data)

    # get all volunteer ids
    volunteer_ids = set([])

    for task_json in tasks['results']:
        t = n.Page(json_obj=task_json)
        volunteers = t.get_list('volunteer')
        for v in volunteers:
            volunteer_ids.add(v)

    # generate schedule for individual volunteers
    for id in volunteer_ids:

        temp_dict = make_volunteer_schedule(id, tasks)
        if temp_dict:
            for error in temp_dict['errors']:
                r['errors'].append(error)
            for warning in temp_dict['warnings']:
                r['warnings'].append(warning)
            for message in temp_dict['messages']:
                r['messages'].append(message)

    return r


def make_volunteer_schedule(id, tasks):

    r = u.get_results_dict()
    logger = logging.getLogger(__name__)

    try:
        v_name = dbfuncs.get_name_from_notion_id(id=id, model=Guest)
    except:
        u.add_to_results(
            r, id, 'e', "Could not match id to name. Update guest db!", "#")
        return r

    # location of our schedule template file
    template = "files/schedule_template.docx"
    doc = Document(template)

    # Title Paragraph
    paragraph = doc.paragraphs[0]
    paragraph.add_run(f"VOLUNTEER SCHEDULE FOR {v_name.upper()}")
    paragraph.style = "Headline 1"
    # As of ...
    today = datetime.datetime.now()
    paragraph = doc.add_paragraph(
        f"(Information as of {today.strftime('%d')} {today.strftime('%b')}.)")
    # make sure these paragraph styles are defined in the template
    paragraph.style = "Text Margins"
    # Blank line
    doc.add_paragraph("")

    seen_days = []

    for task_json in tasks['results']:

        task = n.Page(json_obj=task_json)
        task_volunteer_ids = task.get_list('volunteer')
        if not task_volunteer_ids:
            continue
        if id not in task.get_list('volunteer'):
            continue
        task_time = task.get_date('date')
        if not task_time:
            logger.warning("Skipping task because of missing time.")
            continue

        task_name = task.get_text('task')
        task_event_ids = task.get_list('events')
        task_event_names = []
        if task_event_ids:
            task_event_names = [dbfuncs.get_name_from_notion_id(
                event_id, Event) for event_id in task_event_ids]
        task_venue_ids = task.get_list('venues')
        task_venue_name = dbfuncs.get_name_from_notion_id(
            task_venue_ids[0], Venue) if task_venue_ids else "{unspecified location}"
        task_volunteer_names = [dbfuncs.get_name_from_notion_id(
            volunteer_id, Guest) for volunteer_id in task_volunteer_ids]
        task_notes = task.get_text('notes')

        task_day = u.get_date_string(task_time[0], output='day')
        if task_day not in seen_days:
            seen_days.append(task_day)

            paragraph = doc.add_paragraph(task_day)
            paragraph.style = "Day"
        paragraph = doc.add_paragraph(f"{u.get_date_string(task_time[0], output='time')}-{u.get_date_string(task_time[1], output='time')}"
                                      f": {task_name} at {task_venue_name}")
        paragraph.style = "Event Time + Title"

        if task_event_ids:
            paragraph = doc.add_paragraph(
                f"Event(s): {', '.join(task_event_names)}")
            paragraph.style = "Event Text"

        if len(task_volunteer_names) > 1:
            paragraph = doc.add_paragraph(
                f"Volunteers: {', '.join(task_volunteer_names)}")
            paragraph.style = "Event Text"

        if task_notes:
            paragraph = doc.add_paragraph(f"Notes: {task_notes}")
            paragraph.style = "Event Text"

    # save document
    dropbox_folder = f"{u.get_secret('DROPBOX_FOLDER')}"
    filename = f"{u.get_secret('FESTIVAL_ACRONYM')}_VolunteerSchedule_{v_name}.docx"
    doc.save(f"schedule_temp.docx")

    logger.info(f"locally saved schedule for {v_name}.")
    u.dropbox_upload_local_file(f"schedule_temp.docx",
                                f"{dropbox_folder}/volunteer_schedules/{filename}")
    u.add_to_results(r, f"{v_name}", "m", "Schedule uploaded", "#")
    return r
