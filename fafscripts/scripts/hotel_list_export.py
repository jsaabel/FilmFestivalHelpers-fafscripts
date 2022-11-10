from fafscripts.modules import notion_new as n, utils as u, dbfuncs
from fafscripts.models import Room, Guest
from fafscripts.scripts.database_analyzer import get_results_dict, add_to_results
import datetime
import openpyxl as xl
from docx import Document
from random import randrange
import logging

# https://www.programiz.com/python-programming/datetime/strftime


def main() -> dict:

    logger = logging.getLogger(__name__)

    dropbox_folder = u.get_secret("DROPBOX_FOLDER")

    # results-dict
    r = get_results_dict()

# INITIAL SETUP
    double_rooms = list()
    single_rooms = list()
    special_cases = list()
    now = datetime.datetime.now()
    numbers_for_days = dict()
    decimal_to_day = {
        18: 'Tuesday',
        19: 'Wednesday',
        20: 'Thursday',
        21: 'Friday',
        22: 'Saturday',
        23: 'Sunday',
        24: 'Monday',
    }

    doc = Document("files/schedule_template.docx")
    wb = xl.Workbook()
    sheet = wb.active
    line_count = 1

    sheet.cell(
        line_count, 1).value = f"Hotel list {u.get_secret('FESTIVAL_NAME')}"
    line_count += 1
    sheet.cell(
        line_count, 1).value = f"Print date: {now.strftime('%d/%m/%Y, %H:%M')}"
    line_count += 1
    # sheet.cell(line_count, 1).value = "Name"
    # sheet.cell(line_count, 2).value = "Room"
    # sheet.cell(line_count, 3).value = "Checkin"
    # sheet.cell(line_count, 4).value = "Checkout"
    sheet.cell(line_count, 1).value = "First name"
    sheet.cell(line_count, 2).value = "Last name"
    sheet.cell(line_count, 3).value = "Phone"
    sheet.cell(line_count, 4).value = "Email"
    sheet.cell(line_count, 5).value = "Arrival"
    sheet.cell(line_count, 6).value = "Departure"
    sheet.cell(line_count, 7).value = "Single"
    sheet.cell(line_count, 8).value = "Double"
    sheet.cell(line_count, 9).value = "Sharing with"

    line_count += 1

# CONNECT TO (FILTERED AND SORTED!?) GUEST DB
    data_dict = dict()
    n.add_filter_to_request_dict(request_dict=data_dict, property_name='Room', filter_content=True,
                                 filter_condition="is_not_empty", property_type="relation")
    notion_guests = n.get_db('guests', data_dict=data_dict)
    # item_count_notion = n.get_item_count(notion_guests)

# RETRIEVE RELEVANT PROPERTIES (incl. transformations)
    shared_room_ids = []

    for guest_json in notion_guests['results']:

        g = n.Page(json_obj=guest_json)
        room_id = g.get_list('room')
        room_type = dbfuncs.notion_ids_to_model_props(
            room_id, model=Room)[0]['name']
        if room_type == "None":
            continue
        # implement 'skipping' shared rooms here
        if g.id in shared_room_ids:
            continue

        sharing_with_ids = g.get_list('room-shared-with')
        if sharing_with_ids:
            try:
                sharing_with_names = u.list_to_comma_separated(
                    [dbfuncs.get_name_from_notion_id(id, Guest) for id in sharing_with_ids])
            except:
                logger.error(
                    "Could not match guest to guest id. Try updating 'Guests' internal database.")
                raise RuntimeError
        shared_room_ids += sharing_with_ids

        info_string = str()
        name = g.get_text('name')
        name_split = name.split()
        if len(name_split) == 3:
            name = f"{name_split[2]}, {name_split[0]} {name_split[1]}"
        else:
            name = f"{name_split[1]}, {name_split[0]}"

        first_name = name_split[0]
        last_name = " ".join(name_split[1:]) if len(name_split) > 1 else ""
        phone = g.get_text('phone')
        email = g.get_text('e-mail')

        check_in = g.get_date('checkin')
        if check_in:
            check_in = check_in[0].split("-")
            check_in = datetime.date(day=int(check_in[2]), month=int(
                check_in[1]), year=int(check_in[0]))
        else:
            logger.warning(f"OBS: Date missing for: {name}!")
            check_in = ""
        check_out = g.get_date('checkout')
        if check_out:
            check_out = check_out[0].split("-")
            check_out = datetime.date(day=int(check_out[2]), month=int(
                check_out[1]), year=int(check_out[0]))
        else:
            logger.warning(f"OBS: Date missing for: {name}!")
            check_out = ""

        # compose info string and add to appropriate list
        if check_in and check_out:
            info_string = f"{name}: {check_in.strftime('%a, %d %b')} – {check_out.strftime('%a, %d %B')}"
        else:
            info_string = f"{name}: Dates tbd"

        if room_type == 'Double':
            double_rooms.append(info_string)
        elif room_type == 'Single':
            single_rooms.append(info_string)
        else:
            special_cases.append(f"{info_string}. OBS: {room_type}")

        # EXCEL
        sheet.cell(line_count, 1).value = first_name
        sheet.cell(line_count, 2).value = last_name
        sheet.cell(line_count, 3).value = phone
        sheet.cell(line_count, 4).value = email

        if check_in:
            sheet.cell(line_count, 5).value = check_in.strftime('%d.%m.%y')
        if check_out:
            sheet.cell(line_count, 6).value = check_out.strftime('%d.%m.%y')

        if room_type == 'Single':
            sheet.cell(line_count, 7).value = "x"
        if room_type == 'Double':
            sheet.cell(line_count, 8).value = "x"
        if room_type == 'Double (2 singles)':
            sheet.cell(line_count, 8).value = "x *"
        if sharing_with_ids:
            sheet.cell(line_count, 9).value = sharing_with_names

        line_count += 1

        # rooms/day functionality (new 06.10.2021)
        if check_in and check_out:
            check_in_decimal = int(check_in.strftime('%d'))
            check_out_decimal = int(check_out.strftime('%d'))
            for number in range(check_in_decimal, check_out_decimal):
                if number not in numbers_for_days:
                    numbers_for_days[number] = 1
                else:
                    numbers_for_days[number] += 1

    line_count += 2

    sheet.cell(line_count, 1).value = "* - two singles"
    for key, value in sorted(numbers_for_days.items()):
        add_to_results(
            r, f"{decimal_to_day[key]}", "m", f"{value} room(s) in total.", url="#")


# .DOCX

    p = doc.paragraphs[0]
    p.add_run(f"Room overview {u.get_secret('FESTIVAL_ACRONYM')}")
    p.style = "Day"
    doc.add_paragraph()
    p = doc.add_paragraph(f"Print date: {now.strftime('%d %b %Y, %H:%M')}")
    doc.add_paragraph()

    p = doc.add_paragraph(f'Double rooms ({len(double_rooms)}):')
    p.style = "Event Time + Title"
    for string in sorted(double_rooms):
        doc.add_paragraph(string, style="List Paragraph")

    p = doc.add_paragraph(f'Single rooms ({len(single_rooms)}):')
    p.style = "Event Time + Title"
    for string in sorted(single_rooms):
        doc.add_paragraph(string, style="List Paragraph")

    if special_cases:
        p = doc.add_paragraph(f'Special cases ({len(special_cases)}):')
        p.style = "Event Time + Title"
        for string in sorted(special_cases):
            doc.add_paragraph(string, style="List Paragraph")

# SAVING
    rand = randrange(1000)
    doc.save(f"temp{rand}.docx")
    u.dropbox_upload_local_file(
        f"temp{rand}.docx", f"{dropbox_folder}/hotel_list.docx")
    wb.save(f'temp{rand}.xlsx')
    u.dropbox_upload_local_file(
        f"temp{rand}.xlsx", f"{dropbox_folder}/hotel_list.xlsx")

    return r


if __name__ == "__main__":
    main()
