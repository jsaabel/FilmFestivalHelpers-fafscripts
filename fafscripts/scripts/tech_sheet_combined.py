from fafscripts.modules import notion_new as n, utils as u, dbfuncs
from fafscripts.models import Venue, Guest
from docx import Document
import datetime
from random import randrange
import logging


def main(venue_ids: list):

    # setup + db retrieval
    logger = logging.getLogger(__name__)
    dropbox_folder = u.get_secret("DROPBOX_FOLDER")
    # location of our schedule template file
    template = "files/schedule_template.docx"
    data_dict = dict()
    # n.add_filter_to_request_dict(data_dict, "Venue", "relation", "contains", venue_id)
    # retrieved event db has to be sorted chronologically
    n.add_sorts_to_request_dict(data_dict, 'Time', 'ascending')
    notion_data = n.get_db('events', data_dict=data_dict)
    doc = Document(template)

    # Title Paragraph
    # v_name = dbfuncs.get_name_from_notion_id(venue_id, Venue)
    p = doc.paragraphs[0]
    p.add_run(f"TECH SCHEDULE")
    p.style = "Headline 1"
    # As of ...
    today = datetime.datetime.now()
    p = doc.add_paragraph(
        f"(Information as of {today.strftime('%d')} {today.strftime('%b')}.")
    p.add_run(
        f" Included venues: {', '.join([dbfuncs.get_name_from_notion_id(venue_id, Venue) for venue_id in venue_ids])}.)")
    p.style = "Text Margins"  # make sure these paragraph styles are defined in the template
    # Blank line
    doc.add_paragraph("")

    seen_days = []

    for event_json in notion_data['results']:

        e = n.Page(json_obj=event_json)
        e_venue = e.get_list('venue')
        if not e_venue:
            continue
        if e_venue[0] not in venue_ids:
            continue
        e_name = e.get_text('name')
        e_time = e.get_date('time')
        if not e_time:
            logger.warning("Skipping event because of missing time.")
            continue
        v_name = dbfuncs.get_name_from_notion_id(e_venue[0], Venue)
        e_functionary_ids = e.get_list('functionaries')
        e_descr = e.get_text('description-functionaries')
        e_tech_instr = e.get_text('tech-instructions')
        e_equipment = e.get_text('equipment')
        e_responsibility = e.get_text('responsibility')

        e_day = u.get_date_string(e_time[0], output='day')
        if e_day not in seen_days:
            seen_days.append(e_day)

            if len(seen_days) > 1:
                doc.add_page_break()
                p = doc.add_paragraph()
                p.add_run(f"TECH SCHEDULE")
                p.style = "Headline 1"
                # As of ...
                today = datetime.datetime.now()
                p = doc.add_paragraph(
                    f"(Information as of {today.strftime('%d')} {today.strftime('%b')}.")
                p.add_run(
                    f" Included venues: {', '.join([dbfuncs.get_name_from_notion_id(venue_id, Venue) for venue_id in venue_ids])}.)")
                p.style = "Text Margins"  # make sure these paragraph styles are defined in the template
                # Blank line
                doc.add_paragraph("")
            p = doc.add_paragraph(e_day)
            p.style = "Day"

        p = doc.add_paragraph(f"{u.get_date_string(e_time[0], output='time')}")
        if e_time[1]:
            p.add_run(f"-{u.get_date_string(e_time[1], output='time')}")
        p.add_run(f": {e_name} at {v_name}")
        p.style = "Event Time + Title"

        # NB ADD responsible...
        if e_responsibility:
            p = doc.add_paragraph()
            run = p.add_run(f'Responsible: ')
            run.bold = True
            p.add_run(f"{e_responsibility}")

        if e_functionary_ids:
            e_functionary_names = [dbfuncs.get_name_from_notion_id(
                id, Guest) for id in e_functionary_ids]
            p = doc.add_paragraph()
            run = p.add_run(f'Involved: ')
            run.bold = True
            p.add_run(f"{', '.join(e_functionary_names)}")
            # p.style = "Event Text"

        if e_equipment:
            p = doc.add_paragraph()
            run = p.add_run(f'Equipment: ')
            run.bold = True
            p.add_run(f"{e_equipment}")

        if e_tech_instr:
            p = doc.add_paragraph()
            run = p.add_run(f'Tech instructions: ')
            run.bold = True
            p.add_run(f"{e_tech_instr}")

        if e_descr:
            p = doc.add_paragraph(e_descr)

    # save document
    rand = randrange(1000)
    dropbox_folder = f"{u.get_secret('DROPBOX_FOLDER')}"
    filename = f"{u.get_secret('FESTIVAL_ACRONYM')}_TechSheet_combined.docx"
    doc.save(f"techsheet{rand}.docx")

    # logger.info(f"locally saved schedule for {v_name}.")
    u.dropbox_upload_local_file(f"techsheet{rand}.docx",
                                f"{dropbox_folder}/tech_sheets/{filename}")
