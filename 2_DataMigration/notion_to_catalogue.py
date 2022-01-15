# 30.08.21
# adjusted to changes made in notion 30/31.08

from modules import notion as n, utils as u
import os
from docx import Document  # pip install python-docx

import notion_to_catalogue_schedule, notion_to_catalogue_venues

# GLOBAL VARIABLES
catalogue_folder = "../exports/FAF_Catalogue_21"  # cat. folder name
filetypes = ["txt", "img"]  # sub-folders

# WORKING WITH PAGES
print("\nConnecting to notion database: Catalogue...")
data_dict = dict()
n.add_filter_to_request_dict(data_dict, 'data_source', 'select', 'does_not_equal', '(nn)')
notion_data = n.get_db('catalogue_overview', data_dict=data_dict)
item_count_notion = n.get_item_count(notion_data)

# prompt functionality
categories = list()
for index in range(item_count_notion):
    category_1 = n.get_property(index, 'Category 1', 'select', source=notion_data)
    if category_1 not in categories:
        categories.append(category_1)
categories_sorted = sorted(categories)

print('Found the following categories:')
for category_index, category in enumerate(categories_sorted):
    print(f"{category_index}: {category}")

export_choices = list()
while True:
    export_prompt = input('Choose categories to export (enter indices or q to exit prompt). -> ')
    if export_prompt == "q":
        break
    else:
        export_choices.append(int(export_prompt))


# START LOOP
for index in range(item_count_notion):
    category_1 = n.get_property(index, 'Category 1', 'select', source=notion_data)

    for list_index in export_choices:
        if categories_sorted[list_index] != category_1:  # matching chosen categories
            pass
        else:  # assigning data
            category_1 = n.get_property(index, 'Category 1', 'select', source=notion_data)
            category_2 = n.get_property(index, 'Category 2', 'select', source=notion_data)
            heading = n.get_property(index, 'Heading', 'title', source=notion_data)
            seq = str(n.get_property(index, 'Seq', 'number', source=notion_data)).zfill(2)

            data_source = n.get_property(index, 'data_source', 'select', source=notion_data)

            event_relations = n.get_property(index, 'Events', 'relation', source=notion_data)
            guest_relations = n.get_property(index, 'Guests', 'relation', source=notion_data)
            notion_id = n.get_property(index, property_type='id', source=notion_data)
            info_lines = None
            # if page is linked to an event in event db via relation property, info lines are retrieved/generated from there.
            if event_relations:
                info_lines = n.generate_info_lines(event_relations)

            # assign corresponding notion programme tag
            # OBS this has to correspond EXACTLY to tags used in film db!
            notion_programme_tag = None
            if category_2:
                notion_programme_tag = category_2[3:]  # strip away number in front of programme name [Format: XX Progamme Name]

            # setting folder/save locations and creating them
            if category_2:
                txt_location = f"{catalogue_folder}/{filetypes[0]}/{category_1}/{category_2}"
                img_location = f"{catalogue_folder}/{filetypes[1]}/{category_1}/{category_2}"
            else:
                txt_location = f"{catalogue_folder}/{filetypes[0]}/{category_1}"
                img_location = f"{catalogue_folder}/{filetypes[1]}/{category_1}"
            # making folder locations 'safe' by checking for unsupported characters etc
            txt_location = u.safe_folder_name(txt_location)
            img_location = u.safe_folder_name(img_location)
            folders = [txt_location, img_location]
            for folder in folders:
                try:
                    os.makedirs(folder)
                except FileExistsError:
                    pass

        # EXPORTS

            # dealing with different "data sources"/export modes

            if data_source == 'from file':

                dl_urls = n.get_property(index, 'File', 'files', source=notion_data)
                if not dl_urls:
                    print(f'OBS: Encountered missing file when trying to download {heading}')
                    pass
                else:
                    if ".png" in dl_urls[0]:
                        file_extension = ".png"
                    elif ".pdf" in dl_urls[0]:
                        file_extension = ".pdf"
                    else:
                        file_extension = ".jpg"
                    file_name = f'{seq} {heading}{file_extension}'
                    try:
                        n.download_file(dl_urls[0], target_folder=img_location, target_filename=file_name)
                    except (IndexError, UnboundLocalError):
                        print(f'OBS: Encountered error when trying to download {heading}!')

            elif data_source == 'schedule script':
                notion_to_catalogue_schedule.main(folder_location=txt_location)  # EXECUTE SCHEDULE GENERATOR SCRIPT

            elif data_source == 'venue script':
                notion_to_catalogue_venues.main(folder_location=txt_location)

            elif data_source == 'logo script':  # LOGO DOWNLOAD
                logo_main_folder = f'{img_location}/{seq} {heading}'
                try:
                    os.makedirs(logo_main_folder)
                except FileExistsError:
                    pass
                request_dict = dict()
                n.add_filter_to_request_dict(request_dict, property_name='Catalogue', property_type='checkbox',
                                             filter_condition='equals', filter_content=True)
                notion_supporters = n.get_db('supporters', data_dict=request_dict)
                for supporter in range(n.get_item_count(notion_supporters)):
                    # gather data
                    name = n.get_property(supporter, 'Name', 'title', source=notion_supporters)
                    category = n.get_property(supporter, 'Category', 'select', source=notion_supporters)
                    seq = n.get_property(supporter, 'Seq', 'number', source=notion_supporters)
                    logos_urls = n.get_property(supporter, 'Logos', 'files', source=notion_supporters)
                    # create folders
                    logo_sub_folder_location = f'{logo_main_folder}/{category}/{str(seq).zfill(2)} {name}'
                    try:
                        os.makedirs(logo_sub_folder_location)
                    except FileExistsError:
                        pass
                    # download and rename logo
                    for logo_url in logos_urls:
                        n.download_file(logo_url, target_folder=logo_sub_folder_location)

            else:  # "regular"/non script-related export
                # TEXT
                # initialize doc
                doc = Document()
                # add text
                doc.add_paragraph(heading)
                if info_lines:
                    for info_line in info_lines:
                        paragraph = doc.add_paragraph(info_line)
                    info_lines.clear()

                if data_source == 'from page':  # OBS: text has to be written EXACTLY as specified here.
                    # this allows to incorporate pages that get their content from the respective notion page.
                    # retrieving text blocks and adding them to the document (formatting currently not supported!)
                    text_blocks = n.get_children(notion_id)
                    block_count = n.get_item_count(text_blocks)
                    for block in range(block_count):
                        try:
                            plain_text = n.get_text_from_child(block, source=text_blocks)
                            doc.add_paragraph(plain_text)
                        except IndexError:  # this is necessary to deal with, amongst others, blank lines.
                            doc.add_paragraph()
                else:
                    if event_relations:
                        event_page = n.get_page(event_relations[0])
                        event_text = n.get_property_from_page('Text catalogue', 'text', source=event_page)
                        u.add_text(event_text, doc)
                        event_img = n.get_property_from_page('Pic', 'url', source=event_page)
                        if event_img:
                            u.save_img(img_url=event_img, img_name=f'{seq} {heading}', location=img_location)

                    elif guest_relations:
                        guest_page = n.get_page(guest_relations[0])
                        guest_text = n.get_property_from_page('Bio_eng', 'text', source=guest_page)
                        u.add_text(guest_text, doc)
                        guest_img = n.get_property_from_page('Headshot', 'files', source=guest_page)
                        if guest_img:
                            u.save_first_img(guest_img, f'{img_location}/{seq} {heading}')  # replace with n.download_file?

                    else:
                        pass

                # (add Commissioned Films)
                if notion_programme_tag == 'NBC: Commissioned Film':  # OBS EXACT SPELLING
                    print('Adding commissioned films...')
                    doc.add_paragraph()  # blank line
                    u.add_commissioned_films(doc)  # special function for generating commissioned films document

                # save document
                doc_filename = u.safe_file_name(f"{seq} {heading}.docx")  # making filename 'safe'
                doc.save(f"{txt_location}/{doc_filename}")
                print(f"\tsaved document for: {heading}.")

                # # export image LEGACY
                # if img_url:
                #     img_filename = u.safe_file_name(f"{seq} {heading}")  # making filename 'safe'
                #     u.save_img(img_url, img_filename, img_location)  # this should be more robust to errors

            # generate film files
                if not category_2:  # all films have a 'category 2' property
                    pass
                else:
                    if ('Films' in category_1) and ('Commissioned' not in category_2):  # OBS: Identifying 'Film' pages
                        print(f'Trying to generate films files for {notion_programme_tag}...')
                        data_dict = dict()
                        n.add_filter_to_request_dict(data_dict, 'Programme', 'multi_select', 'contains',
                                                     notion_programme_tag)
                        if notion_programme_tag == "NBC: Children's Film":
                            n.add_sorts_to_request_dict(data_dict, "Seq Children's Film", 'ascending')
                        else:
                            n.add_sorts_to_request_dict(data_dict, 'Seq', 'ascending')
                        films = n.get_db('films', data_dict=data_dict)
                        u.generate_film_files(films, folder_location_text=txt_location, folder_location_images=img_location)
                        # dedicated function for generating film files
