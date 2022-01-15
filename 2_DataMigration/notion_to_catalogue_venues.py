from modules import notion as n
from docx import Document


def main(folder_location: str = "../exports"):
    # retrieve filtered and sorted location db
    data_dict = dict()
    n.add_filter_to_request_dict(data_dict, "Seq", "number", "is_not_empty", True)
    n.add_sorts_to_request_dict(data_dict, 'Seq')
    notion_venues = n.get_db("venues", data_dict=data_dict)
    notion_venues_count = n.get_item_count(notion_venues)
    # compose document
    doc = Document()
    doc.add_paragraph("Venues")
    doc.add_paragraph()
    for i in range(notion_venues_count):
        venue_name = n.get_property(i, 'Name', 'title', source=notion_venues)
        venue_text = n.get_property(i, 'Catalogue Text', 'text', source=notion_venues)
        doc.add_paragraph(venue_name)
        doc.add_paragraph(venue_text)
        doc.add_paragraph()
    # save document
    file_name = "venues.docx"
    doc.save(f'{folder_location}/{file_name}')
    print(f'\t...saved {file_name}.')






