from modules import notion as n, utils as u
from docx import Document

# retrieve event db, applying desired filters
data_dict = dict()
# n.add_filter_to_request_dict(data_dict, property_name='Tech Instructions', property_type='text',
#                              filter_condition='is_not_empty', filter_content=True)
# n.add_filter_to_request_dict(data_dict, property_name='stream?', property_type='checkbox',
#                              filter_condition='equals', filter_content=True)
n.add_sorts_to_request_dict(data_dict, property_name='Time', sorts_direction='ascending')
n.add_filter_to_request_dict(data_dict, property_name='stream?', property_type='checkbox', filter_content=True, filter_condition='equals')
notion_events = n.get_db('events', data_dict=data_dict)
item_count_notion = n.get_item_count(notion_events)

# retrieve venue db
notion_venues = n.get_db('venues')

# retrieve (filtered) guest db
filter_dict = dict()
n.add_filter_to_request_dict(filter_dict, property_name='Events', property_type='relation',
                             filter_condition='is_not_empty', filter_content=True)
notion_guests = n.get_db('guests', data_dict=filter_dict)

# initialize document
doc = Document('../files/Tech Sheet Template.docx')
last_date = ""
# add phone list, general instructions etc? or open template where this is already integrated

# loop through events

for index in range(0, item_count_notion):
    # retrieve data
    event_title = n.get_property(index, 'Name', 'title', source=notion_events)
    try:
        venue_notion_id = n.get_property(index, 'Venue', 'relation', source=notion_events)
        venue_notion_page = n.get_page_from_db(venue_notion_id[0], source=notion_venues)
        event_venue = n.get_property_from_page('Name', 'title', source=venue_notion_page)
    except IndexError:
        pass
    event_time = n.get_property(index, 'Time', 'date', source=notion_events)
    equipment = n.get_property(index, 'Equipment', 'text', source=notion_events)
    tech_instructions = n.get_property(index, 'Tech Instructions', 'text', source=notion_events)
    functionaries_ids = n.get_property(index, 'Functionaries', 'relation', source=notion_events)
    functionary_names = [n.get_name_from_relation(functionary_id, relation_db=notion_guests) for functionary_id in functionaries_ids]
    description_internal = n.get_property(index, 'Description Functionaries', 'text', source=notion_events)
    volunteer_instructions = n.get_property(index, 'Volunteer Instructions', 'text', source=notion_events)
    responsibility = n.get_property(index, 'Responsibility', 'people', source=notion_events)
    comment_anders = n.get_property(index, 'Comment (Anders)', 'text', source=notion_events)

    # filter level 2
    if not 'Kino' in event_venue:
        pass
    elif 'Gamlebyen' in event_venue:
        pass
    else:
        # compose document
        if u.get_date_string(event_time[0], output='day') != last_date:
            last_date = u.get_date_string(event_time[0], output='day')
            p = doc.add_paragraph(u.get_date_string(event_time[0], output='day'))
            p.style = "Title"
        p = doc.add_paragraph(f"{u.get_date_string(event_time[0], output='time')}: "
                                           f"{event_title.upper()}")
        p.style = "Heading 1"
        if event_venue:
            p = doc.add_paragraph(f'{event_venue.upper()}')
        # comment properties in/out as needed
        if responsibility:
            p = doc.add_paragraph()
            run = p.add_run(f'Responsible: ')
            run.bold = True
            p.add_run(f'{responsibility[0]}')
        if equipment:
            p = doc.add_paragraph()
            run = p.add_run(f'Equipment: ')
            run.bold = True
            p.add_run(f'{equipment}')
        # if comment_anders:
        #     p = doc.add_paragraph()
        #     run = p.add_run(f'Comment (Anders): ')
        #     run.bold = True
        #     p.add_run(f'{comment_anders}')
        if tech_instructions:
            p = doc.add_paragraph()
            run = p.add_run(f'Instructions: ')
            run.bold = True
            p.add_run(f'{tech_instructions}')
        # if volunteer_instructions:
        #     p = doc.add_paragraph()
        #     run = p.add_run(f'Volunteer Instructions: ')
        #     run.bold = True
        #     p.add_run(f'{volunteer_instructions}')
        # if functionary_names:
        #     p = doc.add_paragraph()
        #     run = p.add_run(f'Involved: ')
        #     run.bold = True
        #     p.add_run(f'{u.list_to_comma_separated(functionary_names)}')

        doc.add_paragraph()
        doc.add_paragraph(description_internal)
        doc.add_paragraph()



# save document
doc.save('../exports/techsheets.docx')
print('Document saved.')

# to do: separate into days, get filter to work as intended, add 'volunteer tasks' (if exists), determine 'final' format
