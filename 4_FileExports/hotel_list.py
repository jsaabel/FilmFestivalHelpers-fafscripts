"""
(Nov 19, 2021)
Creates list of needed hotel rooms w/ names and other info in word and excel format. Also counts number of rooms needed
for each room type.
Minor possible improvement: List missing information more explicitly
Major possible improvement: Write 'status report' (counts, last update, missing information) to notion db
(possibly running script in background and executing regularly?)
"""
from modules import notion as n
import datetime
import openpyxl as xl
from docx import Document

# https://www.programiz.com/python-programming/datetime/strftime

# INITIAL SETUP
double_rooms = list()
single_rooms = list()
special_cases = list()
now = datetime.datetime.now()
numbers_for_days = dict()
decimal_to_day = {
    19: 'Tuesday',
    20: 'Wednesday',
    21: 'Thursday',
    22: 'Friday',
    23: 'Saturday',
    24: 'Sunday',
    25: 'Monday',
}

doc = Document("../files/schedule_test.docx")
wb = xl.Workbook()
sheet = wb.active
line_count = 1

sheet.cell(line_count, 1).value = "Hotel list Fredrikstad Animation Festival 2021"
line_count += 1
sheet.cell(line_count, 1).value = f"Print date: {now.strftime('%d/%m/%Y, %H:%M')}"
line_count += 1
sheet.cell(line_count, 1).value = "Name"
sheet.cell(line_count, 2).value = "Room"
sheet.cell(line_count, 3).value = "Checkin"
sheet.cell(line_count, 4).value = "Checkout"
line_count += 1

# CONNECT TO (FILTERED AND SORTED!?) GUEST DB
data_dict = dict()
n.add_filter_to_request_dict(request_dict=data_dict, property_name='Room', filter_content=True,
                             filter_condition="is_not_empty", property_type="select")
notion_guests = n.get_db('guests', data_dict=data_dict)
item_count_notion = n.get_item_count(notion_guests)

# RETRIEVE RELEVANT PROPERTIES (incl. transformations)
for index in range(item_count_notion):
    room_type = n.get_property(index, 'Room', 'select', source=notion_guests)
    if room_type == "None":  # Filter
        pass
    else:
        info_string = str()
        name = n.get_property(index, 'Name', 'title', source=notion_guests)
        name_split = name.split()
        if len(name_split) == 3:
            name = f"{name_split[2]}, {name_split[0]} {name_split[1]}"
        else:
            name = f"{name_split[1]}, {name_split[0]}"

        check_in = n.get_property(index, 'Checkin', 'date', source=notion_guests)
        if check_in:
            check_in = check_in[0].split("-")
            check_in = datetime.date(day=int(check_in[2]), month=int(check_in[1]), year=int(check_in[0]))
        else:
            print(f"OBS: Date missing for: {name}!")
        check_out = n.get_property(index, 'Checkout', 'date', source=notion_guests)
        if check_out:
            check_out = check_out[0].split("-")
            check_out = datetime.date(day=int(check_out[2]), month=int(check_out[1]), year=int(check_out[0]))

        # compose info string and add to appropriate list
        if check_in and check_out:
            info_string = f"{name}: {check_in.strftime('%a, %d %b')} – {check_out.strftime('%a, %d %B')}"
        else:
            info_string = f"{name}: Dates tbd"

        if room_type == 'Double':
            double_rooms.append(info_string)
        elif room_type == 'Single':
            single_rooms.append(info_string)
        else:
            special_cases.append(f"{info_string}. OBS: {room_type}")

        # EXCEL
        sheet.cell(line_count, 1).value = name
        sheet.cell(line_count, 2).value = room_type
        if check_in:
            sheet.cell(line_count, 3).value = check_in.strftime('%a, %d.%m')
        if check_out:
            sheet.cell(line_count, 4).value = check_out.strftime('%a, %d.%m')
        line_count += 1

        # rooms/day functionality (new 06.10.2021)
        if check_in and check_out:
            check_in_decimal = int(check_in.strftime('%d'))
            check_out_decimal = int(check_out.strftime('%d'))
            for number in range(check_in_decimal, check_out_decimal + 1):
                if number not in numbers_for_days:
                    numbers_for_days[number] = 1
                else:
                    numbers_for_days[number] += 1

# continue here
for key, value in sorted(numbers_for_days.items()):
    print(f'{decimal_to_day[key]}.: {value}')


# .DOCX

p = doc.paragraphs[0]
p.add_run('Room overview FAF 2021')
p.style = "Day"
doc.add_paragraph()
p = doc.add_paragraph(f"Print date: {now.strftime('%d %b %Y, %H:%M')}")
doc.add_paragraph()

p = doc.add_paragraph(f'Double rooms ({len(double_rooms)}):')
p.style = "Event Time + Title"
for string in sorted(double_rooms):
    doc.add_paragraph(string, style="List Paragraph")

p = doc.add_paragraph(f'Single rooms ({len(single_rooms)}):')
p.style = "Event Time + Title"
for string in sorted(single_rooms):
    doc.add_paragraph(string, style="List Paragraph")

if special_cases:
    p = doc.add_paragraph(f'Special cases ({len(special_cases)}):')
    p.style = "Event Time + Title"
    for string in sorted(special_cases):
        doc.add_paragraph(string, style="List Paragraph")

# SAVING
doc.save("../exports/FAF2021_rooms.docx")
wb.save('../exports/FAF2021_rooms.xlsx')
