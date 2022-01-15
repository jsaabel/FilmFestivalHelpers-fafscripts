# SCHEDULE EXPORT

from modules import notion as n, utils as u
from docx import Document
import os
import datetime
# from docx2pdf import convert  # pip install docx2pdf
# OBS outstanding improvements: Make more robust to missing travel info!


def get_travel_info(guest_name):
    """compile a dictionary with all relevant travel information"""
    # generate dictionary
    # OBS: source has to be titled 'notion_data'
    travel_dict = dict()  # new approach for this part 12.10.2021: Fix data / account for missing info here!
    nn_string = "[not specified]"
    for i in range(item_count_notion):
        if guest_name in n.get_property(i, 'Name', 'title', source=notion_data):
            # try:

            travel_dict["incoming"] = n.get_property(i, 'Incoming', 'text', source=notion_data)
            departure_date_notion = n.get_property(i, 'Departure Incoming', 'date', source=notion_data)
            if departure_date_notion:
                travel_dict["departure_incoming_date"]= u.get_date_string(notion_date=departure_date_notion[0], output='day')
                travel_dict["departure_incoming_time"]= u.get_date_string(notion_date=departure_date_notion[0], output='time')
            else:
                travel_dict["departure_incoming_date"] = nn_string
                travel_dict["departure_incoming_time"] = nn_string
            travel_dict["booking_ref_in"] = n.get_property(i, 'Booking Ref In', 'text', source=notion_data)
            travel_dict["arrival_at"] = n.get_property(i, 'Arrival at', 'select', source=notion_data)
            arrival_date = n.get_property(i, 'Arrival', 'date', source=notion_data)
            if arrival_date:
                travel_dict["arrival_date"] = u.get_date_string(notion_date=arrival_date[0], output='day')
                travel_dict["arrival_time"] = u.get_date_string(notion_date=arrival_date[0], output='time')
            else:
                travel_dict["arrival_date"] = nn_string
                travel_dict["arrival_time"] = nn_string
            travel_dict["transport_incoming"] = n.get_property(i, 'Transport incoming', 'text', source=notion_data)
            travel_dict["outgoing"] = n.get_property(i, 'Outgoing', 'text', source=notion_data)
            travel_dict["booking_ref_out"] = n.get_property(i, 'Booking Ref Out', 'text', source=notion_data)
            travel_dict["departure_outgoing_from"] = n.get_property(i, 'Departure from', 'select', source=notion_data)
            departure_outgoing_date_notion = n.get_property(i, 'Departure Outgoing', 'date', source=notion_data)
            if departure_outgoing_date_notion:
                travel_dict["departure_outgoing_date"] = u.get_date_string(notion_date=departure_outgoing_date_notion[0], output='day')
                travel_dict["departure_outgoing_time"] = u.get_date_string(notion_date=departure_outgoing_date_notion[0], output='time')
            else:
                travel_dict["departure_outgoing_date"] = nn_string
                travel_dict["departure_outgoing_time"] = nn_string
            travel_dict["transport_outgoing"] = n.get_property(i, 'Transport outgoing', 'text', source=notion_data)

            for key, value in travel_dict.items():  # setting all missing dic values to nn_string
                if not value:
                    travel_dict[key] = nn_string

    return travel_dict


template = "../files/schedule_test.docx"  # location of our schedule template file
# set folder location
folder_location = "../exports/schedules" # folder location for export
# retrieve needed databases from notion
# filter: only guests that are assigned to events
# OBS OBS this will be problematic after having implemented improved "attendee/functionary" functionality OBS OBS
# look into multiple ('OR'/'AND') filters?
filter_dict = dict()
n.add_filter_to_request_dict(filter_dict, property_name='Events', property_type='relation',
                             filter_condition='is_not_empty', filter_content=True)
notion_data = n.get_db('guests', data_dict=filter_dict)  # this has to be called notion_data for traveldict func to work
item_count_notion = n.get_item_count(notion_data)
data_dict = dict()
n.add_sorts_to_request_dict(data_dict, 'Time', 'ascending')  # retrieved event db has to be sorted chronologically
events_data = n.get_db('events', data_dict=data_dict)
item_count_events = n.get_item_count(events_data)
notion_venues = n.get_db('venues')

# create subfolder for files
try:
    os.makedirs(folder_location)
except FileExistsError:
    pass

# looping over guests
for index in range(item_count_notion):
    # check if assigned to events, skip guest if not // OBS: Is this even necessary if db is filtered upon retrieval?
    events_1 = n.get_property(index, property_name='Events', property_type='relation', source=notion_data)
    events_2 = n.get_property(index, property_name='Functionary', property_type='relation', source=notion_data)
    if not (events_1 or events_2):
        pass
    else:
        guest_name = n.get_property(index, property_name='Name', property_type='title', source=notion_data)
        guest_id = n.get_property(index, property_type='id', source=notion_data)  # notion unique identifier
        # events are retrieved from event db using relation property
        event_ids_attendee = n.get_property(index, property_name='Events', property_type='relation', source=notion_data)
        event_ids_functionary = n.get_property(index, property_name='Functionary', property_type='relation', source=notion_data)
        event_ids_combined = list(event_ids_attendee)
        event_ids_combined.extend(x for x in event_ids_functionary if x not in event_ids_attendee)
        # composing document
        document = Document(template)  # loading in schedule template
        # Title Paragraph
        paragraph = document.paragraphs[0]
        paragraph.add_run(f"SCHEDULE FOR {guest_name.upper()}")
        paragraph.style = "Headline 1"
        # As of ...
        today = datetime.datetime.now()
        paragraph = document.add_paragraph(f"(Information as of {today.strftime('%d')} {today.strftime('%b')}. "
                                           f"Changes and/or mistakes may occur. Please refer to your ticket(s) for further details.)")
        paragraph.style = "Text Margins"  # make sure these paragraph styles are defined in the template
        # Blank line
        document.add_paragraph("")
        # OPTIONAL: "Intro" paragraph - instructions on eventive, covid etc
        first_name = guest_name.split(" ", 1)[0]
        paragraph = document.add_paragraph(f"Dear {first_name}! We are looking forward to welcoming you to Fredrikstad,"
                                           f" and hope you will find this schedule helpful. Please make sure to read the"
                                           f" description of your events carefully, as rehearsals, tech checks etc."
                                           f" are often scheduled BEFORE the official start time of the event.")
        paragraph.style = "Event Text"

        # TRAVEL INFORMATION (temporatily disabled for easier testing of other functionality)
        # check status of 'Travel?' checkbox
        travel_check = n.get_property(index, 'Travel?', 'checkbox', source=notion_data)
        travel_info = dict()  # getting rid of 'may not be assigned' warning
        if travel_check:

            travel_info = get_travel_info(guest_name)  # return dictionary with travel information
            paragraph = document.add_paragraph("ARRIVAL")
            paragraph.style = "Day"

            paragraph = document.add_paragraph("Departure with ")
            paragraph.style = "Event Text"
            run = paragraph.add_run(f"{travel_info['incoming']} ")
            run.bold = True
            run = paragraph.add_run("at ")
            run.bold = False
            run = paragraph.add_run(f"{travel_info['departure_incoming_time']}")
            run.bold = True
            run = paragraph.add_run(" (local time) on ")
            run.bold = False
            run = paragraph.add_run(f"{travel_info['departure_incoming_date']}. ")
            run.bold = True

            if '[' not in travel_info['booking_ref_in']:
                run = paragraph.add_run(f"Your booking reference: {travel_info['booking_ref_in']}. ")
                run.bold = False

            paragraph = document.add_paragraph(f"Arrival at {travel_info['arrival_at']} at ")
            paragraph.style = "Event Text"
            run = paragraph.add_run(f"{travel_info['arrival_time']}. ")
            run.bold = True

            if '[' not in travel_info['transport_incoming']:
                run = paragraph.add_run(f"Transport to festival: {travel_info['transport_incoming']}. ")
                run.bold = False

        # Split event ids into days
        event_ids_wednesday = list()
        event_ids_thursday = list()
        event_ids_friday = list()
        event_ids_saturday = list()
        event_ids_sunday = list()

        for event_id in event_ids_combined:
            for i2 in range(item_count_events):
                if event_id == events_data['results'][i2]['id']:  # matching event id obtained from relation to event db
                    found_index = i2
                    index_and_id = (found_index, event_id)  # tuple: index in (sorted) json - event id. allows for ordering!

                    date = n.get_property(i2, 'Time', property_type='date', source=events_data)
                    # matching dates and weekdays
                    if '10-20' in date[0]:
                        event_ids_wednesday.append(index_and_id)
                    elif '10-21' in date[0]:
                        event_ids_thursday.append(index_and_id)
                    elif '10-22' in date[0]:
                        event_ids_friday.append(index_and_id)
                    elif '10-23' in date[0]:
                        event_ids_saturday.append(index_and_id)
                    elif '10-24' in date[0]:
                        event_ids_sunday.append(index_and_id)

        weekday_dict = {
            'Wednesday, 20 OCT': sorted(event_ids_wednesday),
            'Thursday, 21 OCT': sorted(event_ids_thursday),
            'Friday, 22 OCT': sorted(event_ids_friday),
            'Saturday, 23 OCT': sorted(event_ids_saturday),
            'Sunday, 24 OCT': sorted(event_ids_sunday),
        }

        for day, ids in weekday_dict.items():
            if len(ids) >= 1:  # -> only add this paragraph if at least one event for the day exists
                paragraph = document.add_paragraph(day)
                paragraph.style = "Day"
                for id in ids:
                    for i3 in range(0, item_count_events):
                        if events_data['results'][i3]['id'] == id[1]:  # assign event info
                            event_title = n.get_property(i3, 'Name', property_type='title', source=events_data)
                            if ' (second screening)' in event_title:
                                event_title = event_title.replace(' (second screening)', '')
                            event_time = n.get_property(i3, 'Time', property_type='date', source=events_data)
                            venue_notion_id = n.get_property(i3, 'Venue', 'relation', source=events_data)
                            venue_notion_page = n.get_page_from_db(venue_notion_id[0], source=notion_venues)
                            event_venue = n.get_property_from_page('Name', 'title', source=venue_notion_page)
                            event_description = n.get_property(i3, 'Description Schedule', source=events_data)
                            event_attendees = n.get_property(i3, 'Attendees', property_type='relation',source=events_data)
                            event_functionaries = n.get_property(i3, 'Functionaries', property_type='relation', source=events_data)
                            event_description_functionaries = n.get_property(i3, 'Description Functionaries', source=events_data)

                            # write to document
                            paragraph = document.add_paragraph(f"{u.get_date_string(event_time[0], output='time')} "
                                                               f"{event_title} at {event_venue}")
                            paragraph.style = "Event Time + Title"
                            if guest_id in event_functionaries:
                                paragraph = document.add_paragraph(event_description_functionaries)
                            else:
                                paragraph = document.add_paragraph(event_description)
                            paragraph.style = "Event Text"

        # outgoing travel information
        if travel_check:
            if not n.get_property(index, 'Departure from', 'select', source=notion_data):
                pass
            else:
                paragraph = document.add_paragraph("DEPARTURE")
                paragraph.style = "Day"
                paragraph = document.add_paragraph("Departure with ")
                paragraph.style = "Event Text"
                run = paragraph.add_run(f"{travel_info['outgoing']} from {travel_info['departure_outgoing_from']}")
                run.bold = True
                run = paragraph.add_run(" at ")
                run.bold = False
                run = paragraph.add_run(f"{travel_info['departure_outgoing_time']}")
                run.bold = True
                run = paragraph.add_run(" on ")
                run.bold = False
                run = paragraph.add_run(f"{travel_info['departure_outgoing_date']}. ")
                run.bold = True

                if '[' not in travel_info['booking_ref_out']:
                    run = paragraph.add_run(f"Your booking reference: {travel_info['booking_ref_out']}. ")
                    run.bold = False

                if '[' not in travel_info['transport_outgoing']:
                    run = paragraph.add_run(f"Transport: {travel_info['transport_outgoing']}. ")

        # Save Document
        filename = f"{folder_location}/FAF 2021_Schedule_{guest_name}.docx"
        try:
            document.save(filename)
            print(f"saved {filename}")
        except OSError:
            pass
        except AttributeError:
            pass
        except TypeError:
            pass

    # Convert (all) word files to pdf
    # convert(folder_location)
