from modules import notion as n, utils as u
from docx import Document

# retrieving a filtered and sorted film db
data_dict = dict()
n.add_filter_to_request_dict(data_dict, 'Programme', 'multi_select', 'contains', 'NBC: Student Film')
n.add_sorts_to_request_dict(data_dict, 'Seq', 'ascending')
data_source = n.get_db('films', data_dict=data_dict)
item_count_notion = n.get_item_count(data_source)

doc = Document()

for index in range(item_count_notion):
    # NB: make sure property names and types match those from database
    film_data = {
        "title": n.get_property(index, "English Title", 'title', source=data_source),
        "title_ov": n.get_property(index, "Original Title", 'text', source=data_source),
        "year": n.get_property(index, "Year", 'select', source=data_source),
        "director": n.get_property(index, "Director", 'text', source=data_source),
        "synopsis": n.get_property(index, "Synopsis", 'text', source=data_source),
        "bio": n.get_property(index, "Bio", 'text', source=data_source),
        "country": u.list_to_comma_separated(n.get_property(index, "Country", 'multi_select', source=data_source)),
        "runtime": n.get_property(index, "Runtime", 'text', source=data_source),  # OBS
        "technique": u.list_to_comma_separated(n.get_property(index, "Technique", 'multi_select', source=data_source)),
        "production": n.get_property(index, "Production", 'text', source=data_source),
        "animation": n.get_property(index, "Animation", 'text', source=data_source),
        "seq": str(index + 2).zfill(2),  # OBS
    }

    # composing and adding a single line for each film in word document
    p = doc.add_paragraph()

    run = p.add_run(film_data.get('title'))
    run.italic = True

    run = p.add_run(f" ({film_data['country']} – {film_data['year']})")
    run.italic = False

    p.add_run(f", regi: {film_data['director']}")

doc.save('../exports/filmlist.docx')  # file location
print('list written.')