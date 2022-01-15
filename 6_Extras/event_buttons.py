"""
(Nov 19, 2021)
Creates a word document containing code snippets for eventive ticket buttons used on website.
Major possible improvement: use wordpress API to add these automatically
"""
from modules import eventive as e, notion as n
from docx import Document

doc = Document()
filter_dict = dict()
n.add_filter_to_request_dict(filter_dict, property_name="eventive_id", property_type="text", filter_condition="is_not_empty", filter_content=True)
notion_events = n.get_db('events', data_dict=filter_dict)
item_count_notion = n.get_item_count(notion_events)

for index in range(item_count_notion):
    event_type = n.get_property(index, 'Type', 'select', source=notion_events)
    if 'screening' in event_type.lower():
        pass
    else:
        event_name = n.get_property(index, 'Name', 'title', source=notion_events)
        eventive_id = n.get_property(index, 'eventive_id', 'text', source=notion_events)
        button_code = f"<div class='eventive-button' data-event='{eventive_id}'></div>"
        button_code = button_code.replace("'", "\"")

        doc.add_paragraph(f'{event_name}:', style='Heading 1')
        doc.add_paragraph(button_code)

doc.save("Test.docx")
