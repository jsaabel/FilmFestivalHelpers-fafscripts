"""
(Nov 19, 2021)
Generates word documents that use wordpress formatting tags for more convenient transfer/uploads of texts to website.
Major possible improvement: Eventually, this should turn into automated upload through wordpress API.
Minor possible improvement: Look at number of blank lines needed to make paragraphs look good.
"""
from docx import Document
from modules import notion as n, utils as u

# Loop through all guests that have a catalogue text
request_dict = dict()
n.add_filter_to_request_dict(request_dict, "Bio_eng", "text", "is_not_empty", True)
n.add_sorts_to_request_dict(request_dict, "Name", "ascending")
notion_guests = n.get_db("guests", data_dict=request_dict)
item_count_notion = n.get_item_count(notion_guests)

doc = Document()

for index in range(item_count_notion):
    name = n.get_property(index, "Name", "title", source=notion_guests)
    doc.add_paragraph(name, style='Heading 1')
    doc.add_paragraph('English bio (prepared for web):', style="Heading 2")
    bio_eng = n.get_property(index, 'Bio_eng', 'text', source=notion_guests)
    u.catalogue_text_to_web_text(bio_eng, doc)
    bio_norw = n.get_property(index, 'Bio_norw', 'text', source=notion_guests)
    if bio_norw:
        doc.add_paragraph('Norwegian bio (prepared for web):', style="Heading 2")
        u.catalogue_text_to_web_text(bio_norw, doc)

doc.save("../exports/cat_to_web_guests.docx")

# Loop through all events that have a catalogue text
request_dict = dict()
n.add_filter_to_request_dict(request_dict, "Text catalogue", "text", "is_not_empty", True)
n.add_sorts_to_request_dict(request_dict, "Name", "ascending")
notion_guests = n.get_db("events", data_dict=request_dict)
item_count_notion = n.get_item_count(notion_guests)

doc = Document()

for index in range(item_count_notion):
    name = n.get_property(index, "Name", "title", source=notion_guests)
    doc.add_paragraph(name, style='Heading 1')
    doc.add_paragraph('English text (prepared for web):', style="Heading 2")
    text_eng = n.get_property(index, 'Text catalogue', 'text', source=notion_guests)
    u.catalogue_text_to_web_text(text_eng, doc)
    text_norw = n.get_property(index, 'Text norw', 'text', source=notion_guests)
    if text_norw:
        doc.add_paragraph('Norwegian text (prepared for web):', style="Heading 2")
        u.catalogue_text_to_web_text(text_norw, doc)

doc.save("../exports/cat_to_web_events.docx")